from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "docs" / "performance" / "jmeter"
STATS_PATH = REPORT_DIR / "html-report" / "statistics.json"
OUT_PATH = REPORT_DIR / "EduMind核心接口JMeter压力测试报告.docx"
SCREENSHOTS = REPORT_DIR / "screenshots"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK_BLUE = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
GRID = "D9E2EC"
WHITE = "FFFFFF"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    # Named override for CJK readability while keeping Calibri as the preset base font.
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGINS):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width_dxa))
    tcw.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="6"):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def apply_table_geometry(table, widths):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}, got {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblpr = tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        row.height = None
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_borders(table)


def set_cell_text(cell, text, *, bold=False, color=None, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_spacing(p, before=16, after=8, line=1.10)
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_paragraph_spacing(p, before=12, after=6, line=1.10)
        run = p.add_run(text)
        set_run_font(run, size=13, bold=True, color=BLUE)
    else:
        set_paragraph_spacing(p, before=8, after=4, line=1.10)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
    return p


def add_body(doc, text, *, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.10)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold, italic=italic, color=color)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=10, line=1.10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9.5, italic=True, color=MUTED)
    return p


def add_key_value_table(doc, rows, widths=(2400, 6960), header=None):
    row_count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=row_count, cols=2)
    apply_table_geometry(table, list(widths))
    r = 0
    if header:
        for c, text in enumerate(header):
            set_cell_shading(table.rows[0].cells[c], LIGHT_GRAY)
            set_cell_text(table.rows[0].cells[c], text, bold=True, color=INK_BLUE)
        r = 1
    for label, value in rows:
        set_cell_text(table.rows[r].cells[0], label, bold=True, color=INK_BLUE)
        set_cell_text(table.rows[r].cells[1], value)
        r += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_data_table(doc, headers, rows, widths, font_size=8.8):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    apply_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(cell, header, bold=True, color=INK_BLUE, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.rows[r_idx].cells[c_idx], value, size=font_size, align=align)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    title_run = p.add_run(title + "：")
    set_run_font(title_run, size=11, bold=True, color=INK_BLUE)
    body_run = p.add_run(text)
    set_run_font(body_run, size=11, color=INK_BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    section.header.is_linked_to_previous = False
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = header_p.add_run("EduMind Performance Test Report")
    set_run_font(hr, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer_p.add_run("Page ")
    set_run_font(fr, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    footer_p.add_run()._r.append(fld_begin)
    footer_p.add_run()._r.append(instr)
    footer_p.add_run()._r.append(fld_end)


def add_title_block(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=4, line=1.10)
    run = p.add_run("EduMind 核心接口 JMeter 压力测试报告")
    set_run_font(run, size=23, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=14, line=1.10)
    run = p.add_run("高频用户访问接口压力测试与 HTML Dashboard 截图归档")
    set_run_font(run, size=13.5, color=MUTED)

    meta_rows = [
        ("测试日期", "2026-07-03"),
        ("测试对象", "EduMind API Gateway http://localhost:5000"),
        ("测试计划", "docs/performance/jmeter/edumind-core-api-pressure.jmx"),
        ("HTML 报告", "docs/performance/jmeter/html-report/index.html"),
        ("原始结果", "docs/performance/jmeter/results/core-api-pressure.jtl"),
    ]
    add_key_value_table(doc, meta_rows, widths=(1700, 7660))


def load_stats():
    with STATS_PATH.open("r", encoding="utf-8") as f:
        data = json.load(f)
    order = [
        "00 POST /auth/login",
        "01 GET /auth/me",
        "02 GET /files?role=student",
        "03 GET /chats?role=student",
        "04 GET /quizzes?role=student",
        "05 GET /flashcards",
        "06 GET /wrongbook/summary",
    ]
    return data, order


def fmt_num(value, digits=2):
    return f"{float(value):.{digits}f}"


def build_doc():
    stats, order = load_stats()
    total = stats["Total"]
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_heading(doc, "1. 测试目标", 1)
    add_body(
        doc,
        "本次只选取用户更可能直接、高频访问的核心接口进行压力测试，重点观察认证态校验、资料库列表、学习记录列表、知识卡片列表和错题本汇总等页面型读取接口在并发访问下的稳定性与响应时间。",
    )
    add_body(
        doc,
        "未纳入本次压测的接口包括大模型生成类接口、文件上传与 RAG 索引重建、删除/修改类接口、WebSocket/SSE 长连接接口。原因是这些接口更容易受外部模型、队列、对象存储或长任务执行时间影响，不适合作为本次高频直接访问接口的性能基线。",
    )

    add_callout(
        doc,
        "测试结论",
        f"30 并发、90 秒持续压测下共 {int(total['sampleCount'])} 个样本，失败数 {int(total['errorCount'])}，错误率 {fmt_num(total['errorPct'])}%，整体吞吐量 {fmt_num(total['throughput'])} requests/s。",
    )

    add_heading(doc, "2. 核心接口范围", 1)
    scope_rows = [
        ("0", "POST /auth/login", "每个虚拟用户登录一次，用于获取 Bearer Token"),
        ("1", "GET /auth/me", "前端登录态校验与用户信息刷新高频调用"),
        ("2", "GET /files?role=student", "学生资料库入口，页面打开与刷新高频调用"),
        ("3", "GET /chats?role=student", "对话历史列表，学习入口高频调用"),
        ("4", "GET /quizzes?role=student", "测验历史列表，学习记录高频调用"),
        ("5", "GET /flashcards", "知识卡片列表，复习页面高频调用"),
        ("6", "GET /wrongbook/summary", "错题本首页汇总，高频读取且有一定聚合计算"),
    ]
    add_data_table(
        doc,
        ["序号", "接口", "纳入原因"],
        scope_rows,
        [700, 3000, 5660],
        font_size=9.2,
    )

    add_heading(doc, "3. 测试环境与策略", 1)
    env_rows = [
        ("后端运行方式", "Docker Compose 启动 API Gateway、identity、learning-content、asset-library、ai-core、media-generation、teaching-content、PostgreSQL、Redis"),
        ("JMeter 运行方式", "Docker 镜像 justb4/jmeter:latest，非 GUI 模式"),
        ("并发线程数", "30"),
        ("升压时间", "20 秒"),
        ("持续时间", "90 秒"),
        ("思考时间", "100-300 ms 随机等待"),
        ("测试账号", "jmeter_user_1 到 jmeter_user_30，每个线程独立账号"),
        ("账号策略说明", "项目登录会使同一用户旧 session 失效，因此使用线程独立账号，避免测试脚本制造伪 401。"),
    ]
    add_key_value_table(doc, env_rows, widths=(2200, 7160))

    add_heading(doc, "4. 测试结果", 1)
    summary_rows = [
        ("总样本数", str(int(total["sampleCount"]))),
        ("失败数", str(int(total["errorCount"]))),
        ("错误率", f"{fmt_num(total['errorPct'])}%"),
        ("整体吞吐量", f"{fmt_num(total['throughput'])} requests/s"),
        ("平均响应时间", f"{fmt_num(total['meanResTime'])} ms"),
        ("中位响应时间", f"{fmt_num(total['medianResTime'])} ms"),
        ("P90 / P95 / P99", f"{fmt_num(total['pct1ResTime'])} / {fmt_num(total['pct2ResTime'])} / {fmt_num(total['pct3ResTime'])} ms"),
        ("最大响应时间", f"{fmt_num(total['maxResTime'])} ms"),
    ]
    add_key_value_table(doc, summary_rows, widths=(2500, 6860))

    metrics_rows = []
    for label in order:
        item = stats[label]
        metrics_rows.append(
            (
                label,
                str(int(item["sampleCount"])),
                f"{fmt_num(item['errorPct'])}%",
                fmt_num(item["meanResTime"]),
                fmt_num(item["pct2ResTime"]),
                fmt_num(item["pct3ResTime"]),
                fmt_num(item["maxResTime"], 0),
                fmt_num(item["throughput"]),
            )
        )
    add_data_table(
        doc,
        ["接口", "样本", "错误率", "平均(ms)", "P95(ms)", "P99(ms)", "最大(ms)", "吞吐(req/s)"],
        metrics_rows,
        [2550, 760, 850, 1030, 1030, 1030, 950, 1160],
        font_size=8.5,
    )

    add_heading(doc, "5. HTML Dashboard 浏览器截图", 1)
    add_body(doc, "以下截图均来自 JMeter 生成的 HTML Dashboard，并通过浏览器打开本地 HTML 文件后截图得到。")

    figures = [
        ("图 1  JMeter Dashboard 总览", "jmeter_dashboard_overview.png"),
        ("图 2  Statistics 与 Errors 区域", "jmeter_statistics_table.png"),
        ("图 3  吞吐量趋势", "jmeter_throughput.png"),
        ("图 4  响应时间百分位", "jmeter_response_times.png"),
        ("图 5  平均响应时间趋势", "jmeter_over_time.png"),
    ]
    for idx, (caption, name) in enumerate(figures):
        if idx in {0, 2, 3, 4}:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(str(SCREENSHOTS / name), width=Inches(6.25))
        add_caption(doc, caption)

    doc.add_page_break()
    add_heading(doc, "6. 结论与建议", 1)
    add_body(
        doc,
        "在 30 并发、90 秒持续压测下，本次选取的核心高频接口全部通过，错误率为 0.00%。这说明当前 Docker Compose 运行形态下的 API Gateway、认证服务、学习内容服务、资料库服务和 PostgreSQL/Redis 基础依赖在该压力级别下保持稳定。",
    )
    add_body(
        doc,
        "响应时间方面，/auth/me 相对较快；列表与汇总类接口平均约 0.9-1.13 秒，P95 约 1.25-1.52 秒。/flashcards 和 /wrongbook/summary 的尾部延迟最高，后续可重点检查认证解析、跨服务调用、PostgreSQL KV 前缀扫描、列表分页和缓存策略。",
    )

    add_heading(doc, "7. 复现命令", 1)
    commands = [
        "docker compose up -d api-gateway",
        "docker run --rm --mount \"type=bind,source=$PWD,target=/work\" justb4/jmeter:latest -n -t /work/docs/performance/jmeter/edumind-core-api-pressure.jmx -l /work/docs/performance/jmeter/results/core-api-pressure.jtl -e -o /work/docs/performance/jmeter/html-report -Jhost=host.docker.internal -Jport=5000 -Jthreads=30 -Jramp=20 -Jduration=90 -JusernamePrefix=jmeter_user_ -Jpassword=jmeter123",
    ]
    for cmd in commands:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=5, line=1.05)
        run = p.add_run(cmd)
        set_run_font(run, size=8.5, name="Courier New", color=INK_BLUE)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
