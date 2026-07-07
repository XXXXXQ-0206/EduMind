from __future__ import annotations

import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\大三下实训\EduMind\EduMind")
OUT = ROOT / "文档" / "EduMind文档" / "6_30第七组提交" / "项目架构会议记录-修订版.docx"

FONT = "Microsoft YaHei"
ACCENT = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)
BLACK = RGBColor(0, 0, 0)
TABLE_WIDTH = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para_runs(paragraph, size=None, bold=None, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, ACCENT, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS.items():
        tag = "w:left" if side == "start" else "w:right" if side == "end" else f"w:{side}"
        element = tc_mar.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="B8C7D9"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def add_table(doc: Document, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade(cell, "F2F4F7")
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            set_para_runs(p, size=9, bold=True, color=BLACK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                set_para_runs(p, size=8.8, color=BLACK)
    doc.add_paragraph()
    return table


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5)


def add_header_footer(doc: Document):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "EduMind 项目架构会议记录"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_runs(header, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "第七组 | 修订版"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_runs(footer, size=8.5, color=MUTED)


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("项目架构会议记录")
    set_run_font(run, size=23, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("修订版：同步当前 EduMind 技术栈与部署架构")
    set_run_font(run, size=13, color=MUTED)

    add_table(
        doc,
        ["事项", "内容"],
        [
            ["会议主题", "EduMind 架构口径修订与文档一致性确认"],
            ["会议时间", "2026-06-30"],
            ["修订日期", date.today().isoformat()],
            ["参会人员", "第七组项目成员"],
            ["会议目标", "统一当前前端、后端、存储、异步任务与部署边界描述，修正文档审核中发现的技术栈偏差。"],
        ],
        [1900, 7460],
    )


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)
    add_header_footer(doc)
    add_title_block(doc)

    doc.add_heading("一、会议背景", level=1)
    p = doc.add_paragraph(
        "本次会议基于文档一致性审核结果，对原项目架构会议记录中与当前实现不一致的技术表述进行修订。"
        "当前 EduMind 已形成以 api-gateway 为统一入口、按业务边界拆分服务、通过共享基础设施支撑 RAG 与生成任务的部署方案。"
    )
    set_para_runs(p, size=10.8)

    doc.add_heading("二、当前架构结论", level=1)
    add_bullet(doc, "前端统一采用 React + Vite + TypeScript 技术栈，构建后由前端容器通过 80 端口对外提供服务。")
    add_bullet(doc, "后端以 Python FastAPI + Uvicorn 为主，api-gateway 作为统一入口，默认对外端口为 5000。")
    add_bullet(doc, "服务按身份认证、学习内容、资源文件、AI 核心、媒体生成、教学内容等边界拆分，并保留本地单体兼容模式。")
    add_bullet(doc, "运行时数据底座采用 PostgreSQL JSONB/KV、pgvector、Redis 和 MinIO/S3，对象文件、向量检索、事件总线和任务租约分别由对应基础设施承担。")
    add_bullet(doc, "生成类任务通过 Celery/Redis generation-worker 执行，前端通过 WebSocket 或 SSE 接收进度与结果。")

    doc.add_heading("三、服务边界与端口", level=1)
    add_table(
        doc,
        ["服务", "职责", "端口/入口"],
        [
            ["api-gateway", "统一API入口、健康检查、转发各业务服务", "5000"],
            ["identity", "账号、会话、鉴权与令牌解析", "5101"],
            ["learning-content", "学习对话、智能笔记、测验、闪卡、考试、辩论与学习计划", "5102"],
            ["asset-library", "文件上传、文件元数据、解析入口、RAG索引状态与音频转写", "5103"],
            ["media-generation", "口语练习、播客生成、外部媒体服务与Bilibili检索", "5104"],
            ["teaching-content", "教案、课件、试卷和教学视频生成", "5105"],
            ["ai-core", "统一LLM与Embedding访问能力", "5106"],
            ["generation-worker", "异步生成任务执行、重试与并发控制", "后台工作进程"],
        ],
        [1700, 5700, 1960],
    )

    doc.add_heading("四、基础设施与数据流", level=1)
    add_table(
        doc,
        ["组件", "当前选型", "会议确认"],
        [
            ["关系与键值数据", "PostgreSQL JSONB/KV", "作为运行时默认数据持久化方案，支撑用户、会话、聊天和业务元数据。"],
            ["向量检索", "pgvector", "承载文件切片向量索引，支撑多文件RAG上下文召回。"],
            ["事件与任务协调", "Redis", "用于事件总线、任务租约、Celery broker/result backend 等运行时能力。"],
            ["对象文件", "MinIO/S3", "用于上传文件、生成音视频、课件等对象存储；本地文件仅作为开发适配或缓存。"],
            ["异步任务", "Celery/Redis", "学习内容、媒体生成和教学内容等耗时生成任务交由后台worker处理。"],
        ],
        [1900, 2300, 5160],
    )

    doc.add_heading("五、修订点确认", level=1)
    add_table(
        doc,
        ["审核问题", "修订后的会议记录口径", "状态"],
        [
            ["运行时存储表述过旧", "明确运行时默认使用 PostgreSQL、pgvector、Redis、MinIO/S3；旧版轻量存储表述不再作为当前架构描述。", "已修正"],
            ["前端技术栈描述不一致", "明确前端为 React + Vite + TypeScript，部署通过 frontend 容器与 Nginx 对外服务。", "已修正"],
            ["缺少关键技术栈关键词", "正文已覆盖 React、pgvector、MinIO、Celery、api-gateway、FastAPI、PostgreSQL、Redis。", "已修正"],
        ],
        [2200, 5600, 1560],
    )

    doc.add_heading("六、会议决议", level=1)
    add_bullet(doc, "后续所有架构类文档以本修订版会议记录的技术口径为准。")
    add_bullet(doc, "涉及部署、数据库、RAG、异步生成和前端架构的文档，应同步更新为当前代码与 docker-compose.yml 中的实际实现。")
    add_bullet(doc, "旧版会议记录保留作历史归档，不再作为当前架构依据。")
    add_bullet(doc, "若后续继续调整服务边界或基础设施，应在会议记录、设计说明书、架构图和接口文档中同步维护。")

    doc.add_heading("七、后续行动项", level=1)
    add_table(
        doc,
        ["行动项", "负责人", "完成标准"],
        [
            ["更新架构图中的服务边界与数据底座", "架构文档负责人", "图中出现 api-gateway、pgvector、MinIO、Celery 等当前组件。"],
            ["复核需求说明书中的技术栈章节", "需求文档负责人", "前端、后端、存储、部署描述与当前项目一致。"],
            ["复核数据库设计说明书", "数据库文档负责人", "数据库对象与PostgreSQL/pgvector运行时口径一致。"],
            ["建立文档更新检查清单", "项目组", "每次功能或技术栈调整后同步检查文档。"],
        ],
        [3300, 1800, 4260],
    )

    doc.add_section(WD_SECTION_START.CONTINUOUS)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
