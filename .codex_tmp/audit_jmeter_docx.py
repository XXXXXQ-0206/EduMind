from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "performance" / "jmeter" / "EduMind核心接口JMeter压力测试报告.docx"

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def main() -> None:
    doc = Document(DOCX)
    with ZipFile(DOCX) as archive:
        names = set(archive.namelist())
        media = sorted(name for name in names if name.startswith("word/media/"))
        footer_xml = archive.read("word/footer1.xml")
        footer = ET.fromstring(footer_xml)

    direct_field_nodes = []
    for paragraph in footer.findall(".//w:p", NS):
        for child in list(paragraph):
            if child.tag in {
                f"{{{NS['w']}}}fldChar",
                f"{{{NS['w']}}}instrText",
            }:
                direct_field_nodes.append(child.tag)

    print(f"docx={DOCX}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"inline_shapes={len(doc.inline_shapes)}")
    print(f"media_files={len(media)}")
    print(f"footer_direct_field_nodes={len(direct_field_nodes)}")
    if len(doc.inline_shapes) != 5:
        raise SystemExit("Expected 5 embedded screenshots.")
    if len(media) != 5:
        raise SystemExit("Expected 5 media files.")
    if direct_field_nodes:
        raise SystemExit("Footer field nodes are not wrapped in runs.")
    if len(doc.tables) < 4:
        raise SystemExit("Expected report tables to be present.")


if __name__ == "__main__":
    main()
