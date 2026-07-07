from __future__ import annotations

import ast
import json
import re
import zipfile
from io import BytesIO
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


ROOT = Path(r"D:\大三下实训\EduMind\EduMind")
DOC_DIR = ROOT / "文档" / "EduMind文档"
OUT_DOCX = DOC_DIR / "EduMind文档一致性审核报告.docx"
OUT_JSON = ROOT / ".codex_tmp" / "edumind_doc_audit.json"
PYTHON = "Python"


@dataclass
class DocAudit:
    path: Path
    kind: str
    readable: bool
    chars: int = 0
    headings: list[str] = field(default_factory=list)
    issues: list[dict[str, str]] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)
    keywords: dict[str, int] = field(default_factory=dict)
    snippets: dict[str, list[str]] = field(default_factory=dict)


CURRENT_FACTS = {
    "backend_runtime": "Python FastAPI 0.115.0 + Uvicorn",
    "frontend_runtime": "React 19 + Vite + TypeScript",
    "service_boundaries": "api-gateway, identity, learning-content, asset-library, ai-core, media-generation, teaching-content, generation-worker",
    "storage": "PostgreSQL JSONB/KV, pgvector, Redis, MinIO/S3 object storage",
    "async_tasks": "Celery/Redis generation worker, SSE/WebSocket live events",
    "api_entry": "Gateway port 5000, service ports 5101-5106, Bilibili bridge 5001",
}

OUTDATED_PATTERNS = [
    ("old_project_name", r"LinkEdu|PageLM|PageLM", "仍出现旧项目名或前身名称。"),
    ("old_backend_node", r"Node\.?js|Express|Koa|NestJS", "后端技术栈疑似仍写为Node/Express；当前主要后端为Python FastAPI。"),
    ("old_storage_sqlite", r"SQLite|sqlite|本地JSON|JSON文件|Keyv|keyv", "运行时存储疑似仍写为SQLite/本地JSON/Keyv；当前默认是PostgreSQL、pgvector、Redis、MinIO/S3，旧存储仅用于迁移或本地适配。"),
    ("old_single_service", r"单体|单体架构|单进程|单一后端", "可能未反映当前Docker Compose中的多服务边界与网关部署。"),
    ("old_ports", r"3000|5173|8000|8080", "可能包含旧端口；当前部署入口主要为80/5000，服务端口为5101-5106，Bilibili桥为5001。"),
    ("old_frontend", r"Vue|Angular|Next\.js", "前端技术栈疑似与当前React/Vite不一致。"),
]

CURRENT_REQUIRED_TERMS = [
    "FastAPI",
    "React",
    "Vite",
    "PostgreSQL",
    "pgvector",
    "Redis",
    "MinIO",
    "Celery",
    "api-gateway",
]


def expr_text(node: ast.AST) -> str:
    try:
        return ast.unparse(node)
    except Exception:
        if isinstance(node, ast.Constant):
            return str(node.value)
        return ""


def route_prefix(tree: ast.Module) -> str:
    for node in tree.body:
        if isinstance(node, ast.Assign) and any(isinstance(t, ast.Name) and t.id == "router" for t in node.targets):
            if isinstance(node.value, ast.Call):
                for kw in node.value.keywords:
                    if kw.arg == "prefix" and isinstance(kw.value, ast.Constant):
                        return str(kw.value.value)
    return ""


def current_routes() -> set[tuple[str, str]]:
    routes: set[tuple[str, str]] = {("GET", "/"), ("GET", "/health")}
    route_dir = ROOT / "backend" / "api" / "routes"
    for file in route_dir.glob("*.py"):
        if file.name == "__init__.py":
            continue
        tree = ast.parse(file.read_text(encoding="utf-8"))
        prefix = route_prefix(tree)
        for node in ast.walk(tree):
            if not isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                continue
            for deco in node.decorator_list:
                if not isinstance(deco, ast.Call) or not isinstance(deco.func, ast.Attribute):
                    continue
                method = deco.func.attr.upper()
                if method not in {"GET", "POST", "PUT", "PATCH", "DELETE", "WEBSOCKET"}:
                    continue
                if deco.args and isinstance(deco.args[0], ast.Constant):
                    path = re.sub(r"/+", "/", f"{prefix}{deco.args[0].value}")
                    if not path.startswith("/"):
                        path = f"/{path}"
                    routes.add((method, path))
    return routes


def read_docx(path: Path) -> tuple[str, list[str]]:
    doc = Document(str(path))
    parts: list[str] = []
    headings: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        parts.append(text)
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            headings.append(text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts), headings[:20]


def read_docx_zip(path: Path) -> tuple[str, list[str]]:
    if not zipfile.is_zipfile(path):
        return "", []
    try:
        return read_docx(path)
    except Exception:
        text = []
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                if name.startswith("word/") and name.endswith(".xml"):
                    xml = zf.read(name).decode("utf-8", errors="ignore")
                    text.extend(re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml))
        return "\n".join(text), []


def read_xlsx(path: Path) -> str:
    wb = load_workbook(BytesIO(path.read_bytes()), data_only=True, read_only=True)
    chunks: list[str] = []
    for ws in wb.worksheets:
        chunks.append(f"[Sheet] {ws.title}")
        for row in ws.iter_rows(values_only=True):
            values = [str(v) for v in row if v not in (None, "")]
            if values:
                chunks.append(" | ".join(values))
    return "\n".join(chunks)


def read_puml(path: Path) -> str:
    for encoding in ("utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except Exception:
            continue
    return path.read_text(errors="ignore")


def printable_runs(text: str) -> list[str]:
    runs = re.findall(r"[\u4e00-\u9fffA-Za-z0-9_:/().,#\-+ ]{4,}", text)
    return [run.strip() for run in runs if len(run.strip()) >= 4]


def read_legacy_doc_binary(path: Path) -> str:
    data = path.read_bytes()
    chunks: list[str] = []
    for encoding in ("utf-16le", "gb18030", "latin1"):
        try:
            decoded = data.decode(encoding, errors="ignore")
        except Exception:
            continue
        chunks.extend(printable_runs(decoded))
    cleaned = []
    seen = set()
    for item in chunks:
        item = re.sub(r"\s+", " ", item)
        if item not in seen:
            seen.add(item)
            cleaned.append(item)
    return "\n".join(cleaned[:2000])


def read_file_text(path: Path) -> tuple[str, list[str], bool, str]:
    if path.name.startswith("~$"):
        return "", [], False, "Word临时锁文件，跳过。"
    suffix = path.suffix.lower()
    if suffix == ".docx" or (suffix == ".doc" and zipfile.is_zipfile(path)):
        text, headings = read_docx_zip(path)
        return text, headings, bool(text.strip()), "OOXML Word文档。"
    if suffix == ".doc":
        try:
            text = read_legacy_doc_binary(path)
        except Exception as exc:
            return "", [], False, f"旧版二进制Word文档抽取失败：{exc}"
        return text, [], bool(text.strip()), "旧版二进制Word文档，仅做近似文本抽取；建议用Word/LibreOffice转换后复核。"
    if suffix in {".xlsx", ".xls"} and zipfile.is_zipfile(path):
        try:
            text = read_xlsx(path)
        except Exception as exc:
            return "", [], False, f"表格文件无法可靠抽取：{exc}"
        return text, [], bool(text.strip()), "OOXML表格文档。"
    if suffix == ".puml":
        return read_puml(path), [], True, "PlantUML源文件。"
    if suffix in {".zip", ".png"}:
        return "", [], False, "归档或图片文件，本轮只记录，不做正文一致性审阅。"
    return "", [], False, "暂不支持的文件类型。"


def snippets_for(text: str, pattern: str) -> list[str]:
    snippets = []
    for match in re.finditer(pattern, text, flags=re.I):
        start = max(0, match.start() - 48)
        end = min(len(text), match.end() + 80)
        snippet = re.sub(r"\s+", " ", text[start:end]).strip()
        if snippet not in snippets:
            snippets.append(snippet)
        if len(snippets) >= 3:
            break
    return snippets


def audit_file(path: Path, routes: set[tuple[str, str]]) -> DocAudit:
    text, headings, readable, note = read_file_text(path)
    audit = DocAudit(path=path, kind=path.suffix.lower() or "folder", readable=readable, chars=len(text), headings=headings)
    audit.notes.append(note)
    if not readable:
        if path.suffix.lower() not in {".zip", ".png"} and not path.name.startswith("~$"):
            audit.issues.append({"severity": "P2", "title": "无法可靠抽取正文", "detail": note})
        return audit

    for key, pattern, message in OUTDATED_PATTERNS:
        matches = re.findall(pattern, text, flags=re.I)
        if matches:
            audit.keywords[key] = len(matches)
            audit.snippets[key] = snippets_for(text, pattern)
            severity = "P1" if key in {"old_backend_node", "old_storage_sqlite"} else "P2"
            audit.issues.append({"severity": severity, "title": message, "detail": f"命中 {len(matches)} 次：{key}"})

    lower = text.lower()
    present = [term for term in CURRENT_REQUIRED_TERMS if term.lower() in lower]
    broad_tech_doc = any(marker in path.name for marker in ["架构", "需求", "设计", "调研"])
    if path.suffix.lower() in {".docx", ".doc"} and audit.chars > 2000 and broad_tech_doc:
        if "技术" in text or "架构" in text or "数据库" in text or "系统" in text:
            missing = [term for term in CURRENT_REQUIRED_TERMS if term not in present]
            if len(missing) >= 5:
                audit.issues.append(
                    {
                        "severity": "P2",
                        "title": "缺少当前技术栈关键字",
                        "detail": "未明显覆盖：" + "、".join(missing[:8]),
                    }
                )

    if "接口文档" in path.name:
        doc_paths = set(re.findall(r"/[A-Za-z0-9_\-/{}]+", text))
        current_paths = {item[1] for item in routes if item[1] != "/"}
        missing = sorted(current_paths - doc_paths)
        if missing:
            audit.issues.append({"severity": "P1", "title": "接口文档缺少当前代码路由", "detail": "缺少：" + "、".join(missing[:12])})
        audit.notes.append(f"当前代码路由 {len(routes)} 个，文档提取路径 {len(doc_paths)} 个。")

    return audit


def run_audit() -> tuple[list[DocAudit], set[tuple[str, str]]]:
    routes = current_routes()
    audits: list[DocAudit] = []
    for path in sorted(DOC_DIR.rglob("*")):
        if path.is_file() and path.resolve() != OUT_DOCX.resolve():
            audits.append(audit_file(path, routes))
    OUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(
        json.dumps(
            [
                {
                    "path": str(item.path),
                    "kind": item.kind,
                    "readable": item.readable,
                    "chars": item.chars,
                    "issues": item.issues,
                    "notes": item.notes,
                    "snippets": item.snippets,
                }
                for item in audits
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return audits, routes


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    font = "Microsoft YaHei"
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 77, 120)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text: str, header: bool = False) -> None:
    cell.text = text
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_run_font(run, size=8.6, bold=header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell(table.rows[0].cells[idx], header, True)
        shade(table.rows[0].cells[idx], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value)
            if idx == 0 and value in {"P0", "P1", "P2"}:
                shade(cells[idx], {"P0": "F4CCCC", "P1": "FCE5CD", "P2": "FFF2CC"}[value])
    doc.add_paragraph()


def write_report(audits: list[DocAudit], routes: set[tuple[str, str]]) -> None:
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("EduMind文档一致性审核报告")
    set_run_font(run, size=22, bold=True, color=RGBColor(0, 0, 0))
    p = doc.add_paragraph(f"生成日期：{date.today().isoformat()}；审核对象：{DOC_DIR}")
    for run in p.runs:
        set_run_font(run, size=10, color=RGBColor(89, 89, 89))

    issue_count = sum(len(item.issues) for item in audits)
    readable_count = sum(1 for item in audits if item.readable)
    doc.add_heading("一、结论摘要", level=1)
    summary = [
        f"共扫描 {len(audits)} 个文件，其中 {readable_count} 个可抽取正文，{issue_count} 条潜在偏差。",
        "当前代码事实：后端以Python FastAPI为主，Docker Compose按api-gateway、identity、learning-content、asset-library、ai-core、media-generation、teaching-content和generation-worker拆分运行。",
        "当前数据与基础设施事实：运行时默认使用PostgreSQL/JSONB、pgvector、Redis、MinIO/S3，并通过Celery/Redis承载生成类异步任务。",
        "旧版二进制 .doc 文件缺少本机转换工具，仅做近似文本抽取；建议安装LibreOffice或用Word另存为 .docx 后再复核排版与全文。",
    ]
    for item in summary:
        doc.add_paragraph(item)

    doc.add_heading("二、当前项目事实源", level=1)
    add_table(
        doc,
        ["项目事实", "当前值"],
        [[key, value] for key, value in CURRENT_FACTS.items()] + [["当前路由数量", str(len(routes))]],
    )

    doc.add_heading("三、重点问题", level=1)
    rows = []
    for audit in audits:
        for issue in audit.issues:
            rows.append(
                [
                    issue["severity"],
                    str(audit.path.relative_to(DOC_DIR)),
                    issue["title"],
                    issue["detail"][:180],
                ]
            )
    if not rows:
        rows = [["-", "-", "未发现明显偏差", ""]]
    add_table(doc, ["级别", "文件", "问题", "说明"], rows)

    doc.add_heading("四、逐文件审核", level=1)
    rows = []
    for audit in audits:
        rows.append(
            [
                str(audit.path.relative_to(DOC_DIR)),
                audit.kind,
                "是" if audit.readable else "否",
                str(audit.chars),
                f"{len(audit.issues)}条",
                "；".join(audit.notes)[:180],
            ]
        )
    add_table(doc, ["文件", "类型", "已读正文", "字符数", "问题数", "备注"], rows)

    doc.add_heading("五、命中片段示例", level=1)
    for audit in audits:
        if not audit.snippets:
            continue
        doc.add_heading(str(audit.path.relative_to(DOC_DIR)), level=2)
        for key, snippets in audit.snippets.items():
            doc.add_paragraph(f"{key}:")
            for snippet in snippets:
                doc.add_paragraph(snippet, style=None)

    doc.save(OUT_DOCX)


def main() -> None:
    audits, routes = run_audit()
    write_report(audits, routes)
    print(str(OUT_DOCX))
    print(str(OUT_JSON))


if __name__ == "__main__":
    main()
