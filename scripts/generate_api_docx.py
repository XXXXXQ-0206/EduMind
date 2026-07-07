from __future__ import annotations

import ast
import copy
import json
import re
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional reference extraction
    PdfReader = None


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
ROUTES_DIR = BACKEND / "api" / "routes"
REFERENCE_PDF = Path(r"D:\大三下实训\甄选资料2026\甄选资料2026\接口文档\订单相关接口文档.pdf")
OUTPUT_DIR = ROOT / "文档" / "EduMind文档"
OUTPUT_DOCX = OUTPUT_DIR / "EduMind接口文档.docx"

ACCENT = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)
BLACK = RGBColor(0, 0, 0)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BORDER = "B8C7D9"
FONT = "Microsoft YaHei"
MONO = "Consolas"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


@dataclass
class ModelField:
    name: str
    type_text: str
    required: bool
    default: str = ""


@dataclass
class ModelInfo:
    name: str
    fields: list[ModelField] = field(default_factory=list)


@dataclass
class ParamInfo:
    name: str
    type_text: str
    source: str
    required: bool
    model: str = ""


@dataclass
class EndpointInfo:
    module: str
    file: Path
    line: int
    tag: str
    boundary: str
    boundary_desc: str
    method: str
    path: str
    function: str
    docstring: str
    request_model: str = ""
    request_content_type: str = "application/x-www-form-urlencoded"
    response_content_type: str = "application/json"
    auth: str = "否"
    params: list[ParamInfo] = field(default_factory=list)
    response_examples: list[dict[str, Any]] = field(default_factory=list)
    status_codes: set[int] = field(default_factory=set)
    response_model: str = ""


SERVICE_BOUNDARIES: dict[str, dict[str, str]] = {
    "auth": {
        "boundary": "identity",
        "description": "账号、会话、鉴权和令牌解析。",
    },
    "chat": {
        "boundary": "learning-content",
        "description": "学习对话、智能笔记、测验、闪卡、辩论、考试与学习计划。",
    },
    "notes": {
        "boundary": "learning-content",
        "description": "学习对话、智能笔记、测验、闪卡、辩论、考试与学习计划。",
    },
    "quiz": {
        "boundary": "learning-content",
        "description": "学习对话、智能笔记、测验、闪卡、辩论、考试与学习计划。",
    },
    "flashcards": {
        "boundary": "learning-content",
        "description": "学习对话、智能笔记、测验、闪卡、辩论、考试与学习计划。",
    },
    "companion": {
        "boundary": "learning-content",
        "description": "学习对话、智能笔记、测验、闪卡、辩论、考试与学习计划。",
    },
    "exam": {
        "boundary": "learning-content",
        "description": "学习对话、智能笔记、测验、闪卡、辩论、考试与学习计划。",
    },
    "debate": {
        "boundary": "learning-content",
        "description": "学习对话、智能笔记、测验、闪卡、辩论、考试与学习计划。",
    },
    "tasks": {
        "boundary": "learning-content",
        "description": "学习对话、智能笔记、测验、闪卡、辩论、考试与学习计划。",
    },
    "planner": {
        "boundary": "learning-content",
        "description": "学习对话、智能笔记、测验、闪卡、辩论、考试与学习计划。",
    },
    "files": {
        "boundary": "asset-library",
        "description": "文件库、用户文件元数据、RAG索引和音频转写。",
    },
    "transcriber": {
        "boundary": "asset-library",
        "description": "文件库、用户文件元数据、RAG索引和音频转写。",
    },
    "ai_core": {
        "boundary": "ai-core",
        "description": "内部LLM调用和流式生成能力。",
    },
    "speaking": {
        "boundary": "media-generation",
        "description": "口语练习、音频、外部媒体服务和Bilibili检索。",
    },
    "podcast": {
        "boundary": "media-generation",
        "description": "口语练习、音频、外部媒体服务和Bilibili检索。",
    },
    "bilibili": {
        "boundary": "media-generation",
        "description": "口语练习、音频、外部媒体服务和Bilibili检索。",
    },
    "slides": {
        "boundary": "teaching-content",
        "description": "教案、课件、试卷和教学视频生成。",
    },
    "lesson_plan": {
        "boundary": "teaching-content",
        "description": "教案、课件、试卷和教学视频生成。",
    },
    "paper": {
        "boundary": "teaching-content",
        "description": "教案、课件、试卷和教学视频生成。",
    },
    "teaching_video": {
        "boundary": "teaching-content",
        "description": "教案、课件、试卷和教学视频生成。",
    },
}

MODULE_NAMES = {
    "auth": "用户认证接口",
    "chat": "学习对话接口",
    "notes": "智能笔记接口",
    "quiz": "测验与错题本接口",
    "flashcards": "闪卡接口",
    "companion": "文档伴读接口",
    "exam": "考试生成接口",
    "debate": "辩论训练接口",
    "tasks": "任务事件接口",
    "planner": "学习计划接口",
    "files": "文件库接口",
    "transcriber": "音频转写接口",
    "ai_core": "AI核心内部接口",
    "speaking": "口语练习接口",
    "podcast": "播客生成接口",
    "bilibili": "Bilibili检索接口",
    "slides": "课件生成接口",
    "lesson_plan": "教案生成接口",
    "paper": "试卷生成接口",
    "teaching_video": "教学视频接口",
    "health": "基础健康检查接口",
}

BOUNDARY_NAMES = {
    "identity": "身份认证服务",
    "learning-content": "学习内容服务",
    "asset-library": "资源文件服务",
    "ai-core": "AI核心服务",
    "media-generation": "媒体生成服务",
    "teaching-content": "教学内容服务",
    "platform": "平台基础服务",
}

FIELD_DESCRIPTIONS = {
    "username": "用户名",
    "password": "密码",
    "old_password": "旧密码",
    "new_password": "新密码",
    "confirm_password": "确认新密码",
    "authorization": "Authorization请求头，格式为Bearer token",
    "x_internal_service_token": "内部服务调用令牌",
    "token": "登录会话令牌",
    "topic": "生成主题",
    "title": "标题",
    "prompt": "提示词或用户输入",
    "message": "消息内容",
    "messages": "对话消息列表",
    "role": "端侧角色，可取student或teacher",
    "file": "上传文件列表",
    "files": "文件列表",
    "file_id": "文件ID",
    "materialIds": "关联素材ID列表",
    "materials": "素材列表",
    "includeMaterials": "是否使用文件库素材",
    "count": "生成数量",
    "difficulty": "难度等级",
    "query": "检索关键词",
    "maxChunks": "最大召回片段数",
    "maxChars": "最大上下文字符数",
    "chat_id": "对话ID",
    "note_id": "笔记ID",
    "quiz_id": "测验ID",
    "deck_id": "卡组ID",
    "card_id": "闪卡ID",
    "lesson_plan_id": "教案ID",
    "paper_id": "试卷ID",
    "slide_id": "课件ID",
    "video_id": "教学视频ID",
    "pid": "播客ID",
    "filename": "文件名",
    "record_id": "口语记录ID",
    "task_id": "任务ID",
    "kind": "任务类型",
    "debate_id": "辩论ID",
    "questionId": "题目ID",
    "selectedAnswer": "选择的选项下标",
    "correct": "是否答对",
    "question": "题目内容",
    "selectedOption": "选择项文本",
    "correctOption": "正确选项文本",
    "explanation": "解析说明",
    "at": "作答时间戳",
    "answers": "作答记录列表",
    "ok": "业务处理结果",
    "error": "错误信息",
    "data": "业务数据",
    "user": "用户信息",
    "users": "用户列表",
    "stream": "WebSocket流地址",
    "events": "SSE事件地址",
    "status": "状态",
    "context": "RAG检索上下文",
    "chunks": "检索片段",
    "failedFiles": "索引失败文件",
}

ENDPOINT_SUMMARIES = {
    ("POST", "/auth/register"): "用户注册",
    ("POST", "/auth/login"): "用户登录",
    ("POST", "/auth/logout"): "退出登录",
    ("GET", "/auth/me"): "获取当前用户信息",
    ("POST", "/auth/internal/resolve"): "内部服务解析用户令牌",
    ("POST", "/auth/change-password"): "修改密码",
    ("POST", "/auth/delete-account"): "删除账号",
    ("POST", "/chat"): "创建学习对话",
    ("WEBSOCKET", "/ws/chat"): "订阅对话生成过程",
    ("GET", "/chats"): "查询对话列表",
    ("GET", "/chats/{chat_id}"): "获取对话详情",
    ("DELETE", "/chats/{chat_id}"): "删除对话",
    ("POST", "/smartnotes"): "创建智能笔记",
    ("WEBSOCKET", "/ws/smartnotes"): "订阅智能笔记生成过程",
    ("GET", "/smartnotes"): "查询智能笔记列表",
    ("GET", "/smartnotes/{note_id}"): "获取智能笔记详情",
    ("DELETE", "/smartnotes/{note_id}"): "删除智能笔记",
    ("POST", "/quiz"): "创建测验生成任务",
    ("WEBSOCKET", "/ws/quiz"): "订阅测验生成过程",
    ("GET", "/quizzes"): "查询测验列表",
    ("GET", "/quizzes/{quiz_id}"): "获取测验详情",
    ("DELETE", "/quizzes/{quiz_id}"): "删除测验",
    ("GET", "/quizzes/{quiz_id}/attempts"): "获取测验作答记录",
    ("POST", "/quizzes/{quiz_id}/attempts"): "保存测验作答记录",
    ("POST", "/quizzes/{quiz_id}/attempts/answer"): "更新单题作答记录",
    ("GET", "/wrongbook/summary"): "获取错题本统计",
    ("POST", "/wrongbook/report"): "生成错题本报告",
    ("POST", "/wrongbook/weak-points"): "生成薄弱知识点",
    ("POST", "/flashcards"): "创建手动闪卡",
    ("GET", "/flashcards"): "查询闪卡列表",
    ("DELETE", "/flashcards/{card_id}"): "删除闪卡",
    ("POST", "/flashcards/decks"): "生成知识点闪卡卡组",
    ("GET", "/flashcards/decks"): "查询知识点卡组列表",
    ("GET", "/flashcards/decks/{deck_id}"): "获取知识点卡组详情",
    ("DELETE", "/flashcards/decks/{deck_id}"): "删除知识点卡组",
    ("POST", "/api/companion/ask"): "文档伴读问答",
    ("GET", "/exams"): "查询考试列表",
    ("POST", "/exam"): "创建考试生成任务",
    ("WEBSOCKET", "/ws/exams"): "订阅考试生成过程",
    ("POST", "/debate/start"): "开始辩论",
    ("GET", "/debate/{debate_id}"): "获取辩论详情",
    ("GET", "/debates"): "查询辩论列表",
    ("DELETE", "/debate/{debate_id}"): "删除辩论",
    ("POST", "/debate/{debate_id}/argue"): "提交辩论发言",
    ("POST", "/debate/{debate_id}/surrender"): "结束辩论",
    ("POST", "/debate/{debate_id}/analyze"): "生成辩论分析",
    ("WEBSOCKET", "/ws/debate"): "订阅辩论对话流",
    ("WEBSOCKET", "/ws/debate/analyze"): "订阅辩论分析流",
    ("GET", "/tasks/{kind}/{task_id}/events"): "订阅任务SSE事件",
    ("POST", "/tasks/ingest"): "从文本导入学习任务",
    ("GET", "/tasks"): "查询学习任务列表",
    ("GET", "/tasks/{task_id}"): "获取学习任务详情",
    ("POST", "/tasks"): "创建学习任务",
    ("PATCH", "/tasks/{task_id}"): "更新学习任务",
    ("DELETE", "/tasks/{task_id}"): "删除学习任务",
    ("POST", "/tasks/{task_id}/plan"): "生成任务学习计划",
    ("POST", "/planner/weekly"): "生成周学习计划",
    ("POST", "/tasks/{task_id}/materials"): "生成任务学习材料",
    ("POST", "/tasks/{task_id}/files"): "上传任务附件",
    ("DELETE", "/tasks/{task_id}/files/{file_id}"): "删除任务附件",
    ("WEBSOCKET", "/ws/planner"): "订阅学习计划事件",
    ("GET", "/files"): "查询文件列表",
    ("POST", "/files"): "上传文件",
    ("POST", "/files/rag/search"): "检索文件库RAG上下文",
    ("GET", "/files/{file_id}/rag"): "获取文件RAG索引状态",
    ("POST", "/files/{file_id}/rag/index"): "重建文件RAG索引",
    ("DELETE", "/files/{file_id}"): "删除文件",
    ("POST", "/transcriber"): "音频转写",
    ("POST", "/ai/internal/invoke"): "内部LLM调用",
    ("POST", "/ai/internal/invoke/stream"): "内部LLM流式调用",
    ("POST", "/speaking/generate"): "生成口语练习题",
    ("POST", "/speaking/upload"): "上传口语音频",
    ("POST", "/speaking/evaluate"): "评测口语音频",
    ("GET", "/speaking/history"): "查询口语历史",
    ("POST", "/speaking/history"): "保存口语历史",
    ("GET", "/speaking/history/{record_id}"): "获取口语记录详情",
    ("DELETE", "/speaking/history/{record_id}"): "删除口语记录",
    ("POST", "/podcast"): "创建播客生成任务",
    ("WEBSOCKET", "/ws/podcast"): "订阅播客生成过程",
    ("GET", "/podcast/download/{pid}/{filename}"): "下载播客音频",
    ("GET", "/podcasts"): "查询播客列表",
    ("GET", "/podcasts/{pid}"): "获取播客详情",
    ("DELETE", "/podcasts/{pid}"): "删除播客",
    ("GET", "/api/bilibili/search"): "搜索Bilibili视频",
    ("POST", "/slides/generate"): "生成课件",
    ("GET", "/slides/{slide_id}"): "获取课件详情",
    ("GET", "/slides/{slide_id}/download"): "下载课件文件",
    ("GET", "/slides"): "查询课件列表",
    ("POST", "/lesson-plan"): "生成教案",
    ("GET", "/lesson-plans"): "查询教案列表",
    ("GET", "/lesson-plans/{lesson_plan_id}"): "获取教案详情",
    ("DELETE", "/lesson-plans/{lesson_plan_id}"): "删除教案",
    ("GET", "/lesson-plans/{lesson_plan_id}/pdf"): "导出教案PDF",
    ("POST", "/paper"): "生成试卷",
    ("WEBSOCKET", "/ws/paper"): "订阅试卷生成过程",
    ("GET", "/papers"): "查询试卷列表",
    ("GET", "/papers/{paper_id}"): "获取试卷详情",
    ("DELETE", "/papers/{paper_id}"): "删除试卷",
    ("GET", "/papers/{paper_id}/pdf"): "导出试卷PDF",
    ("POST", "/teaching-video"): "创建教学视频生成任务",
    ("WEBSOCKET", "/ws/teaching-video"): "订阅教学视频生成过程",
    ("GET", "/teaching-videos"): "查询教学视频列表",
    ("GET", "/teaching-videos/{video_id}"): "获取教学视频详情",
    ("GET", "/teaching-videos/{video_id}/video"): "播放教学视频",
    ("DELETE", "/teaching-videos/{video_id}"): "删除教学视频",
    ("GET", "/"): "根路径健康检查",
    ("GET", "/health"): "服务健康检查",
}

FUNCTION_VERBS = {
    "create": "创建",
    "start": "开始",
    "list": "查询列表",
    "get": "获取详情",
    "delete": "删除",
    "download": "下载",
    "export": "导出",
    "upload": "上传",
    "generate": "生成",
    "save": "保存",
    "update": "更新",
    "upsert": "更新",
    "rebuild": "重建",
    "search": "检索",
    "stream": "流式订阅",
    "transcribe": "转写",
    "evaluate": "评测",
}


def expr_to_text(node: ast.AST | None) -> str:
    if node is None:
        return ""
    if isinstance(node, ast.Name):
        return node.id
    if isinstance(node, ast.Constant):
        return json.dumps(node.value, ensure_ascii=False) if isinstance(node.value, str) else repr(node.value)
    if isinstance(node, ast.Attribute):
        base = expr_to_text(node.value)
        return f"{base}.{node.attr}" if base else node.attr
    if isinstance(node, ast.Subscript):
        return f"{expr_to_text(node.value)}[{expr_to_text(node.slice)}]"
    if isinstance(node, ast.Tuple):
        return ", ".join(expr_to_text(item) for item in node.elts)
    if isinstance(node, ast.List):
        return f"List[{', '.join(expr_to_text(item) for item in node.elts)}]"
    if isinstance(node, ast.BinOp) and isinstance(node.op, ast.BitOr):
        return f"{expr_to_text(node.left)} | {expr_to_text(node.right)}"
    if isinstance(node, ast.Call):
        return expr_to_text(node.func)
    try:
        return ast.unparse(node)
    except Exception:
        return ""


def default_to_text(node: ast.AST | None) -> str:
    if node is None:
        return ""
    if isinstance(node, ast.Constant):
        return "" if node.value is None else str(node.value)
    if isinstance(node, ast.List):
        return "[]"
    if isinstance(node, ast.Dict):
        return "{}"
    if isinstance(node, ast.Call):
        for keyword in node.keywords:
            if keyword.arg == "default":
                return default_to_text(keyword.value)
        if node.args:
            return default_to_text(node.args[0])
        return ""
    return expr_to_text(node)


def strip_optional(type_text: str) -> tuple[str, bool]:
    text = type_text.strip()
    optional = False
    patterns = [
        (r"^Optional\[(.*)\]$", 1),
        (r"^(.*)\s*\|\s*None$", 1),
        (r"^None\s*\|\s*(.*)$", 1),
    ]
    for pattern, group in patterns:
        match = re.match(pattern, text)
        if match:
            text = match.group(group).strip()
            optional = True
            break
    return text, optional


def normalize_type(type_text: str) -> str:
    text = type_text.replace("typing.", "").replace("dict[", "Dict[").replace("list[", "List[")
    replacements = {
        "str": "string",
        "int": "integer",
        "float": "number",
        "bool": "boolean",
        "Any": "object",
        "UploadFile": "file",
        "Request": "object",
        "WebSocket": "websocket",
    }
    text = strip_optional(text)[0]
    return replacements.get(text, text)


def class_is_basemodel(node: ast.ClassDef) -> bool:
    return any(expr_to_text(base).endswith("BaseModel") for base in node.bases)


def parse_models(tree: ast.Module) -> dict[str, ModelInfo]:
    models: dict[str, ModelInfo] = {}
    for node in tree.body:
        if not isinstance(node, ast.ClassDef) or not class_is_basemodel(node):
            continue
        model = ModelInfo(name=node.name)
        for item in node.body:
            if isinstance(item, ast.AnnAssign) and isinstance(item.target, ast.Name):
                type_text = expr_to_text(item.annotation)
                _, optional = strip_optional(type_text)
                default = default_to_text(item.value)
                required = item.value is None and not optional
                model.fields.append(
                    ModelField(
                        name=item.target.id,
                        type_text=normalize_type(type_text),
                        required=required,
                        default=default,
                    )
                )
        models[model.name] = model
    return models


def get_router_prefix(tree: ast.Module) -> str:
    for node in tree.body:
        if not isinstance(node, ast.Assign):
            continue
        if not any(isinstance(target, ast.Name) and target.id == "router" for target in node.targets):
            continue
        if isinstance(node.value, ast.Call) and expr_to_text(node.value.func).endswith("APIRouter"):
            for kw in node.value.keywords:
                if kw.arg == "prefix" and isinstance(kw.value, ast.Constant):
                    return str(kw.value.value)
    return ""


def route_from_decorator(deco: ast.AST) -> tuple[str, str, dict[str, str]] | None:
    if not isinstance(deco, ast.Call):
        return None
    func = deco.func
    if not isinstance(func, ast.Attribute):
        return None
    if expr_to_text(func.value) != "router":
        return None
    method = func.attr.upper()
    if method not in {"GET", "POST", "PUT", "PATCH", "DELETE", "WEBSOCKET"}:
        return None
    if not deco.args or not isinstance(deco.args[0], ast.Constant):
        return None
    path = str(deco.args[0].value)
    kwargs: dict[str, str] = {}
    for keyword in deco.keywords:
        if keyword.arg:
            kwargs[keyword.arg] = expr_to_text(keyword.value)
    return method, path, kwargs


def call_name(node: ast.AST | None) -> str:
    if isinstance(node, ast.Call):
        return expr_to_text(node.func)
    return ""


def is_call_default(default: ast.AST | None, names: set[str]) -> bool:
    name = call_name(default)
    return any(name.endswith(item) for item in names)


def path_param_names(path: str) -> set[str]:
    return set(re.findall(r"\{([^{}]+)\}", path))


def parameter_info(
    arg: ast.arg,
    default: ast.AST | None,
    path_names: set[str],
    models: dict[str, ModelInfo],
) -> tuple[ParamInfo | None, str | None, str]:
    name = arg.arg
    type_text = expr_to_text(arg.annotation)
    normalized = normalize_type(type_text)

    if name in {"self", "request", "background_tasks"} or normalized in {"Request", "BackgroundTasks", "websocket"}:
        return None, None, ""
    if is_call_default(default, {"Depends"}):
        depends_text = expr_to_text(default)
        auth = "是" if "require_auth" in depends_text else ""
        return None, None, auth
    if name in path_names:
        return ParamInfo(name=name, type_text=normalized, source="path", required=True), None, ""
    if is_call_default(default, {"Header"}):
        return ParamInfo(name=name, type_text=normalized, source="header", required="default=None" not in expr_to_text(default)), None, ""
    if is_call_default(default, {"File"}):
        return ParamInfo(name=name, type_text=normalized, source="form-data/file", required=True), None, ""
    if is_call_default(default, {"Form"}):
        return ParamInfo(name=name, type_text=normalized, source="form-data", required="None" not in expr_to_text(default)), None, ""
    if normalized in models:
        return ParamInfo(name=name, type_text=normalized, source="body", required=True, model=normalized), normalized, ""
    if normalized.startswith("Dict[") or normalized == "dict":
        return ParamInfo(name=name, type_text="object", source="body", required=default is None), None, ""

    _, optional = strip_optional(type_text)
    required = default is None and not optional
    return ParamInfo(name=name, type_text=normalized, source="query", required=required), None, ""


def literal_to_example(node: ast.AST) -> Any:
    if isinstance(node, ast.Constant):
        return node.value
    if isinstance(node, ast.Dict):
        result = {}
        for key, value in zip(node.keys, node.values):
            if isinstance(key, ast.Constant) and isinstance(key.value, str):
                result[key.value] = literal_to_example(value)
        return result
    if isinstance(node, ast.List):
        return [literal_to_example(item) for item in node.elts[:1]]
    if isinstance(node, ast.Tuple):
        return [literal_to_example(item) for item in node.elts[:1]]
    if isinstance(node, ast.Name):
        if node.id in {"True", "False"}:
            return node.id == "True"
        if node.id == "None":
            return None
        return sample_value(node.id)
    if isinstance(node, ast.Call):
        func = expr_to_text(node.func)
        if func.endswith("JSONResponse"):
            for keyword in node.keywords:
                if keyword.arg == "content":
                    return literal_to_example(keyword.value)
        return sample_value(func)
    if isinstance(node, ast.JoinedStr):
        return "string"
    return sample_value(expr_to_text(node))


def collect_response_examples(function: ast.AsyncFunctionDef | ast.FunctionDef) -> tuple[list[dict[str, Any]], set[int]]:
    examples: list[dict[str, Any]] = []
    status_codes: set[int] = set()
    for node in ast.walk(function):
        if isinstance(node, ast.Return):
            value = literal_to_example(node.value) if node.value is not None else None
            if isinstance(value, dict):
                examples.append(value)
        if isinstance(node, ast.Call) and expr_to_text(node.func).endswith("JSONResponse"):
            content = None
            for keyword in node.keywords:
                if keyword.arg == "content":
                    content = literal_to_example(keyword.value)
                if keyword.arg == "status_code" and isinstance(keyword.value, ast.Constant):
                    try:
                        status_codes.add(int(keyword.value.value))
                    except Exception:
                        pass
            if isinstance(content, dict):
                examples.append(content)
    deduped: list[dict[str, Any]] = []
    seen = set()
    for item in examples:
        compact = json.dumps(item, ensure_ascii=False, sort_keys=True)
        if compact not in seen:
            seen.add(compact)
            deduped.append(item)
    return deduped[:3], status_codes


def has_call(function: ast.AsyncFunctionDef | ast.FunctionDef, name: str) -> bool:
    for node in ast.walk(function):
        if isinstance(node, ast.Call) and name in expr_to_text(node.func):
            return True
    return False


def parse_endpoints() -> tuple[list[EndpointInfo], dict[str, dict[str, ModelInfo]]]:
    endpoints: list[EndpointInfo] = []
    module_models: dict[str, dict[str, ModelInfo]] = {}
    for path in sorted(ROUTES_DIR.glob("*.py")):
        if path.name == "__init__.py":
            continue
        module = path.stem
        source = path.read_text(encoding="utf-8")
        tree = ast.parse(source)
        models = parse_models(tree)
        module_models[module] = models
        prefix = get_router_prefix(tree)
        service = SERVICE_BOUNDARIES.get(module, {"boundary": module, "description": ""})
        for node in tree.body:
            if not isinstance(node, (ast.AsyncFunctionDef, ast.FunctionDef)):
                continue
            for deco in node.decorator_list:
                parsed = route_from_decorator(deco)
                if not parsed:
                    continue
                method, local_path, kwargs = parsed
                full_path = normalize_path(prefix + local_path)
                path_names = path_param_names(full_path)
                params: list[ParamInfo] = []
                request_model = ""
                auth = ""
                defaults = [None] * (len(node.args.args) - len(node.args.defaults)) + list(node.args.defaults)
                for arg, default in zip(node.args.args, defaults):
                    param, model_name, auth_marker = parameter_info(arg, default, path_names, models)
                    if param:
                        params.append(param)
                    if model_name and not request_model:
                        request_model = model_name
                    if auth_marker:
                        auth = auth_marker
                if method == "WEBSOCKET" and has_call(node, "require_websocket_auth"):
                    auth = "是"
                if has_call(node, "require_auth"):
                    auth = "是"
                if any(param.source == "header" and param.name == "authorization" for param in params):
                    auth = "是"
                examples, status_codes = collect_response_examples(node)
                decorator_status = kwargs.get("status_code")
                if decorator_status and decorator_status.isdigit():
                    status_codes.add(int(decorator_status))
                endpoint = EndpointInfo(
                    module=module,
                    file=path,
                    line=node.lineno,
                    tag=module,
                    boundary=service["boundary"],
                    boundary_desc=service["description"],
                    method=method,
                    path=full_path,
                    function=node.name,
                    docstring=clean_docstring(ast.get_docstring(node) or ""),
                    request_model=request_model,
                    request_content_type=infer_request_content_type(method, params, request_model),
                    response_content_type=infer_response_content_type(method, full_path, node.name),
                    auth=auth or "否",
                    params=params,
                    response_examples=examples,
                    status_codes=status_codes or ({101} if method == "WEBSOCKET" else {200}),
                    response_model=kwargs.get("response_model", ""),
                )
                endpoints.append(endpoint)

    endpoints.extend(health_endpoints())
    return endpoints, module_models


def health_endpoints() -> list[EndpointInfo]:
    return [
        EndpointInfo(
            module="health",
            file=BACKEND / "core" / "app_factory.py",
            line=67,
            tag="health",
            boundary="platform",
            boundary_desc="应用根路径与健康状态检查。",
            method="GET",
            path="/",
            function="root",
            docstring="Root health check.",
            request_content_type="application/x-www-form-urlencoded",
            response_examples=[{"status": "ok", "message": "EduMind Python Backend is running", "version": "1.0.0"}],
            status_codes={200},
        ),
        EndpointInfo(
            module="health",
            file=BACKEND / "core" / "app_factory.py",
            line=78,
            tag="health",
            boundary="platform",
            boundary_desc="应用根路径与健康状态检查。",
            method="GET",
            path="/health",
            function="health",
            docstring="Health check endpoint.",
            request_content_type="application/x-www-form-urlencoded",
            response_examples=[{"status": "healthy", "llm_provider": "provider", "emb_provider": "provider", "service_boundaries": []}],
            status_codes={200},
        ),
    ]


def normalize_path(value: str) -> str:
    value = re.sub(r"/+", "/", value)
    return value if value.startswith("/") else f"/{value}"


def clean_docstring(value: str) -> str:
    value = " ".join(value.split())
    mojibake_markers = ("鐨", "鍒", "娴", "璺", "涓")
    return "" if any(marker in value for marker in mojibake_markers) else value


def infer_request_content_type(method: str, params: list[ParamInfo], request_model: str) -> str:
    if method == "WEBSOCKET":
        return "WebSocket连接参数"
    if any(param.source.startswith("form-data") for param in params):
        return "multipart/form-data"
    if request_model or any(param.source == "body" for param in params):
        return "application/json"
    return "application/x-www-form-urlencoded"


def infer_response_content_type(method: str, path: str, function_name: str) -> str:
    if method == "WEBSOCKET":
        return "WebSocket JSON消息"
    if path.endswith("/events"):
        return "text/event-stream"
    if path.endswith("/pdf"):
        return "application/pdf"
    if "/download/" in path:
        return "audio/*"
    if path.endswith("/download"):
        return "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    if path.endswith("/video"):
        return "video/*"
    if "stream" in function_name and "internal/invoke/stream" in path:
        return "text/event-stream"
    return "application/json"


def endpoint_summary(endpoint: EndpointInfo) -> str:
    override = ENDPOINT_SUMMARIES.get((endpoint.method, endpoint.path))
    if override:
        return override
    if endpoint.docstring:
        return endpoint.docstring
    parts = endpoint.function.replace("_handler", "").split("_")
    verb = FUNCTION_VERBS.get(parts[0], parts[0])
    noun = " ".join(parts[1:]).replace("ws", "WebSocket")
    if noun:
        return f"{verb}{noun}"
    return endpoint.function


def endpoint_description(endpoint: EndpointInfo) -> str:
    summary = endpoint_summary(endpoint)
    if endpoint.method == "WEBSOCKET":
        return f"{summary}，用于前端实时接收生成状态、增量内容或最终结果。"
    if endpoint.response_content_type == "text/event-stream":
        return f"{summary}，通过SSE持续返回任务事件。"
    if "internal" in endpoint.path:
        return f"{summary}，主要供服务内部调用。"
    return f"{summary}。"


def sample_value(type_text: str, field_name: str = "") -> Any:
    name = field_name or type_text
    lower = name.lower()
    if lower.endswith("id") or lower in {"count", "maxchunks", "maxchars", "questionid", "selectedanswer"}:
        return 1
    if lower in {"correct", "includematerials", "ok"}:
        return True
    if "list" in lower or type_text.startswith("List[") or type_text.startswith("array"):
        return []
    if type_text in {"integer", "int"}:
        return 1
    if type_text in {"number", "float"}:
        return 1.0
    if type_text in {"boolean", "bool"}:
        return True
    if type_text == "file":
        return "<binary>"
    if lower == "role":
        return "student"
    if lower == "difficulty":
        return "medium"
    if lower == "topic":
        return "函数与导数"
    if lower == "username":
        return "student001"
    if lower == "password":
        return "password123"
    return "string"


def request_example(endpoint: EndpointInfo, models: dict[str, ModelInfo]) -> Any:
    if endpoint.method == "WEBSOCKET":
        return f"{endpoint.path}?token=<BearerToken>"
    if endpoint.request_model and endpoint.request_model in models:
        return model_example(endpoint.request_model, models)
    body_params = [param for param in endpoint.params if param.source == "body"]
    if body_params:
        return {param.name: sample_value(param.type_text, param.name) for param in body_params}
    form_params = [param for param in endpoint.params if param.source.startswith("form-data")]
    if form_params:
        return {param.name: sample_value(param.type_text, param.name) for param in form_params}
    query_params = [param for param in endpoint.params if param.source == "query"]
    if query_params:
        query = "&".join(f"{param.name}={sample_value(param.type_text, param.name)}" for param in query_params)
        return f"{endpoint.path}?{query}"
    return "暂无"


def model_example(model_name: str, models: dict[str, ModelInfo], seen: set[str] | None = None) -> dict[str, Any]:
    seen = seen or set()
    if model_name in seen:
        return {}
    seen.add(model_name)
    model = models.get(model_name)
    if not model:
        return {}
    result = {}
    for item in model.fields:
        type_text = item.type_text
        inner_match = re.match(r"(List|list)\[(.*)\]", type_text)
        if inner_match:
            inner = inner_match.group(2).strip()
            if inner in models:
                result[item.name] = [model_example(inner, models, seen)]
            else:
                result[item.name] = [sample_value(inner, item.name)]
        elif type_text in models:
            result[item.name] = model_example(type_text, models, seen)
        else:
            result[item.name] = sample_value(type_text, item.name)
    return result


def response_example(endpoint: EndpointInfo) -> Any:
    if endpoint.method == "WEBSOCKET":
        return {"type": "event", "status": "running", "data": {}}
    if endpoint.response_examples:
        preferred = sorted(endpoint.response_examples, key=lambda item: (not item.get("ok", False), -len(item)))[0]
        return simplify_example(preferred)
    if endpoint.response_content_type == "text/event-stream":
        return "event: message\ndata: {\"ok\": true, \"status\": \"running\"}\n\n"
    if endpoint.response_content_type != "application/json":
        return f"<{endpoint.response_content_type} binary>"
    return {"ok": True}


def simplify_example(value: Any) -> Any:
    if isinstance(value, dict):
        result = {}
        for key, item in value.items():
            result[key] = simplify_example(item)
        return result
    if isinstance(value, list):
        return [simplify_example(value[0])] if value else []
    if isinstance(value, str):
        return value if len(value) <= 80 else "string"
    if value is None:
        return None
    return value


def flatten_response_params(example: Any) -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []

    def walk(prefix: str, item: Any, level: int = 0) -> None:
        display = ("  " * level) + prefix
        typ = value_type(item)
        rows.append((display, FIELD_DESCRIPTIONS.get(prefix, ""), typ, ""))
        if isinstance(item, dict):
            for key, value in item.items():
                walk(key, value, level + 1)
        elif isinstance(item, list) and item and isinstance(item[0], dict):
            for key, value in item[0].items():
                walk(key, value, level + 1)

    if isinstance(example, dict):
        for key, value in example.items():
            walk(key, value)
    else:
        rows.append(("data", "响应内容", value_type(example), ""))
    return rows[:24]


def value_type(item: Any) -> str:
    if isinstance(item, bool):
        return "boolean"
    if isinstance(item, int):
        return "integer"
    if isinstance(item, float):
        return "number"
    if isinstance(item, str):
        return "string"
    if isinstance(item, list):
        return "array"
    if isinstance(item, dict):
        return "object"
    if item is None:
        return "object"
    return type(item).__name__


def reference_pdf_info() -> tuple[int, str]:
    if not REFERENCE_PDF.exists() or PdfReader is None:
        return 0, ""
    reader = PdfReader(str(REFERENCE_PDF))
    text = ""
    for page in reader.pages[:2]:
        text += page.extract_text() or ""
    return len(reader.pages), " ".join(text.split())[:260]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        tag = "w:left" if side == "start" else "w:right" if side == "end" else f"w:{side}"
        element = tc_mar.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, font: str = FONT, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, ACCENT, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title_block(doc: Document, endpoint_count: int, pdf_pages: int) -> None:
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.text = "EduMind API接口文档"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_font(p, size=8.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("EduMind API接口文档")
    set_run_font(run, size=23, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("依据后端FastAPI路由源码静态整理，结构参考“订单相关接口文档.pdf”。")
    set_run_font(run, size=11.5, color=MUTED)

    rows = [
        ("文档版本", "V1.0"),
        ("生成日期", date.today().isoformat()),
        ("项目", "EduMind AI-powered learning platform backend"),
        ("覆盖范围", f"{endpoint_count}个HTTP/WebSocket/SSE接口"),
        ("参考文档", f"订单相关接口文档.pdf（{pdf_pages or '未读取页数'}页）"),
        ("基础路径", "本地默认后端服务根路径，接口路径以源码路由为准"),
    ]
    add_key_value_table(doc, rows)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [1900, 7460])
    set_table_borders(table)
    for idx, (key, value) in enumerate(rows):
        cells = table.rows[idx].cells
        cells[0].text = key
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, color=BLACK, bold=cell is cells[0])
    doc.add_paragraph()


def add_text_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], LIGHT_BLUE)
        for p in hdr[idx].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=8.5, color=BLACK, bold=True)
    for row_data in rows or [["暂无", "", "", "", ""][: len(headers)]]:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data[: len(headers)]):
            cells[idx].text = str(value)
            if idx == 0 and value in {"GET", "POST", "DELETE", "PATCH", "PUT", "WEBSOCKET"}:
                set_cell_shading(cells[idx], LIGHT_GRAY)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, font=MONO if idx in {0, 1} else FONT, size=8.2, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, content: Any) -> None:
    if content == "暂无":
        p = doc.add_paragraph("暂无")
        p.paragraph_format.space_after = Pt(4)
        set_paragraph_font(p, size=9.2, color=MUTED)
        return
    text = content if isinstance(content, str) else json.dumps(content, ensure_ascii=False, indent=2)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, "D9E2EC")
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F7F9FB")
    p = cell.paragraphs[0]
    p.text = text[:2600]
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, font=MONO, size=8.0, color=BLACK)


def add_intro(doc: Document) -> None:
    doc.add_heading("简介", level=1)
    for text in [
        "本文档按服务边界和接口模块整理EduMind后端公开接口。每个接口项沿用参考文档中的核心结构：接口地址、请求方式、请求数据类型、响应数据类型、接口描述、请求示例、请求参数、响应状态、响应参数和响应示例。",
        "鉴权约定：除注册、登录、健康检查和部分内部/公开资源接口外，业务接口通常需要Authorization请求头，格式为Bearer token。WebSocket接口通常通过token查询参数完成鉴权。",
        "响应约定：当前代码以JSON对象为主，常见成功字段为ok=true；生成型接口通常先返回任务ID、WebSocket流地址或SSE事件地址，再通过实时通道推送进度和结果。",
    ]:
        p = doc.add_paragraph(text)
        set_paragraph_font(p, size=10.5)


def add_overview(doc: Document, endpoints: list[EndpointInfo]) -> None:
    doc.add_heading("接口总览", level=1)
    rows = []
    for endpoint in endpoints:
        rows.append(
            [
                BOUNDARY_NAMES.get(endpoint.boundary, endpoint.boundary),
                MODULE_NAMES.get(endpoint.module, endpoint.module),
                endpoint.method,
                endpoint.path,
                endpoint_summary(endpoint),
                endpoint.auth,
            ]
        )
    add_text_table(
        doc,
        ["服务边界", "模块", "方法", "接口地址", "接口说明", "鉴权"],
        rows,
        [1500, 1400, 900, 2700, 2200, 660],
    )


def add_endpoint(doc: Document, endpoint: EndpointInfo, models: dict[str, ModelInfo]) -> None:
    doc.add_heading(endpoint_summary(endpoint), level=3)
    add_key_value_table(
        doc,
        [
            ("接口地址", endpoint.path),
            ("请求方式", endpoint.method),
            ("请求数据类型", endpoint.request_content_type),
            ("响应数据类型", endpoint.response_content_type),
            ("是否需要鉴权", endpoint.auth),
            ("源码位置", f"{endpoint.file.relative_to(ROOT)}:{endpoint.line}"),
        ],
    )
    p = doc.add_paragraph()
    p.add_run("接口描述：")
    p.add_run(endpoint_description(endpoint))
    set_paragraph_font(p, size=9.8)

    p = doc.add_paragraph("请求示例：")
    set_paragraph_font(p, size=9.5, bold=True)
    add_code_block(doc, request_example(endpoint, models))

    param_rows = request_param_rows(endpoint, models)
    add_text_table(
        doc,
        ["参数名称", "参数说明", "请求类型", "是否必须", "数据类型", "schema"],
        param_rows,
        [1700, 2600, 1200, 900, 1600, 1360],
    )

    status_rows = [[str(code), status_description(code, endpoint), endpoint.response_model or ""] for code in sorted(endpoint.status_codes)]
    add_text_table(doc, ["状态码", "说明", "schema"], status_rows, [1400, 5000, 2960])

    example = response_example(endpoint)
    response_rows = flatten_response_params(example)
    add_text_table(
        doc,
        ["参数名称", "参数说明", "类型", "schema"],
        [list(row) for row in response_rows],
        [2200, 3800, 1600, 1760],
    )

    p = doc.add_paragraph("响应示例：")
    set_paragraph_font(p, size=9.5, bold=True)
    add_code_block(doc, example)


def request_param_rows(endpoint: EndpointInfo, models: dict[str, ModelInfo]) -> list[list[str]]:
    rows: list[list[str]] = []
    for param in endpoint.params:
        if param.source == "body" and param.model and param.model in models:
            rows.append(
                [
                    param.name,
                    "请求参数实体",
                    "body",
                    "true",
                    param.model,
                    param.model,
                ]
            )
            for field_item in models[param.model].fields:
                rows.append(
                    [
                        f"  {field_item.name}",
                        FIELD_DESCRIPTIONS.get(field_item.name, ""),
                        "",
                        "true" if field_item.required else "false",
                        field_item.type_text,
                        "",
                    ]
                )
        else:
            rows.append(
                [
                    param.name,
                    FIELD_DESCRIPTIONS.get(param.name, ""),
                    param.source,
                    "true" if param.required else "false",
                    param.type_text,
                    param.model,
                ]
            )
    return rows or [["暂无", "无请求参数", "", "", "", ""]]


def status_description(code: int, endpoint: EndpointInfo) -> str:
    if code == 101:
        return "Switching Protocols"
    if code == 202:
        return "Accepted，任务已创建"
    if code == 400:
        return "请求参数错误"
    if code == 401:
        return "未登录或令牌失效"
    if code == 403:
        return "无权限"
    if code == 404:
        return "资源不存在"
    if code >= 500:
        return "服务端处理失败"
    return "OK"


def group_endpoints(endpoints: list[EndpointInfo]) -> dict[str, dict[str, list[EndpointInfo]]]:
    grouped: dict[str, dict[str, list[EndpointInfo]]] = {}
    for endpoint in endpoints:
        grouped.setdefault(endpoint.boundary, {}).setdefault(endpoint.module, []).append(endpoint)
    return grouped


def boundary_order(name: str) -> int:
    order = ["platform", "identity", "learning-content", "asset-library", "ai-core", "media-generation", "teaching-content"]
    return order.index(name) if name in order else 999


def build_docx() -> None:
    endpoints, module_models = parse_endpoints()
    endpoints.sort(key=lambda item: (boundary_order(item.boundary), item.module, item.path, item.method))
    pdf_pages, _ = reference_pdf_info()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_title_block(doc, len(endpoints), pdf_pages)
    add_intro(doc)
    add_overview(doc, endpoints)

    grouped = group_endpoints(endpoints)
    for boundary in sorted(grouped, key=boundary_order):
        doc.add_heading(BOUNDARY_NAMES.get(boundary, boundary), level=1)
        desc = next((ep.boundary_desc for modules in grouped[boundary].values() for ep in modules if ep.boundary_desc), "")
        if desc:
            p = doc.add_paragraph(desc)
            set_paragraph_font(p, size=10.3, color=MUTED)
        for module in sorted(grouped[boundary]):
            doc.add_heading(MODULE_NAMES.get(module, module), level=2)
            for endpoint in grouped[boundary][module]:
                add_endpoint(doc, endpoint, module_models.get(endpoint.module, {}))

    section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.text = "EduMind API接口文档"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(footer, size=8.5, color=MUTED)

    doc.save(OUTPUT_DOCX)
    print(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_docx()
