"""DOCX template support for lesson plan generation."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Set, Tuple


PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}|\[\[\s*([^\[\]]+?)\s*\]\]")

KEY_ALIASES: dict[str, str] = {
    "title": "title",
    "topic": "title",
    "lesson_title": "title",
    "lessonplan_title": "title",
    "课题": "title",
    "课题名称": "title",
    "教案标题": "title",
    "标题": "title",
    "generated_at": "generated_at",
    "date": "generated_at",
    "日期": "generated_at",
    "生成时间": "generated_at",
    "teaching_goals": "teaching_goals",
    "goals": "teaching_goals",
    "教学目标": "teaching_goals",
    "三维目标": "teaching_goals",
    "目标": "teaching_goals",
    "knowledge": "knowledge",
    "知识与技能": "knowledge",
    "知识技能": "knowledge",
    "process_goal": "process_goal",
    "过程与方法": "process_goal",
    "emotion": "emotion",
    "情感态度与价值观": "emotion",
    "情感目标": "emotion",
    "key_points": "key_points",
    "重点": "key_points",
    "教学重点": "key_points",
    "重难点": "key_and_difficult_points",
    "教学重难点": "key_and_difficult_points",
    "difficult_points": "difficult_points",
    "难点": "difficult_points",
    "教学难点": "difficult_points",
    "preparation": "preparation",
    "教学准备": "preparation",
    "准备": "preparation",
    "process": "process",
    "teaching_process": "process",
    "教学过程": "process",
    "课堂流程": "process",
    "教学流程": "process",
    "教学环节": "process",
    "homework": "homework",
    "作业": "homework",
    "作业设计": "homework",
    "课后作业": "homework",
    "教学内容": "title",
    "授课内容": "title",
    "课程内容": "title",
    "课程名称": "title",
    "课程题目": "title",
    "板书": "blackboard_design",
    "板书设计": "blackboard_design",
}

CORE_CONTENT_KEYS = {
    "teaching_goals",
    "knowledge",
    "process_goal",
    "emotion",
    "key_points",
    "key_and_difficult_points",
    "difficult_points",
    "preparation",
    "process",
    "homework",
    "blackboard_design",
}

LABEL_PRIORITY = [
    ("key_and_difficult_points", ("教学重难点", "重难点")),
    ("teaching_goals", ("教学目标", "三维目标", "学习目标")),
    ("knowledge", ("知识与技能", "知识技能")),
    ("process_goal", ("过程与方法",)),
    ("emotion", ("情感态度与价值观", "情感目标")),
    ("key_points", ("教学重点", "重点")),
    ("difficult_points", ("教学难点", "难点")),
    ("preparation", ("教学准备", "课前准备", "准备")),
    ("process", ("教学过程", "教学流程", "课堂流程", "教学环节")),
    ("homework", ("作业设计", "课后作业", "作业")),
    ("blackboard_design", ("板书设计", "板书")),
    ("title", ("课题名称", "课题", "教学内容", "授课内容", "课程名称", "标题")),
]

PROCESS_GRID_DETAIL_ROLES = {"time", "teacher", "student", "intent", "content"}

PROCESS_PHASE_KEYWORDS = {
    "homework": ("作业", "课后"),
    "lead_in": ("导入", "引入", "热身", "复习"),
    "practice": ("练习", "巩固", "训练", "检测", "反馈"),
    "summary": ("小结", "总结", "归纳"),
    "explore": ("合作", "探究", "讨论", "实验", "活动"),
    "new": ("新授", "新知", "讲授", "讲解", "知识建构"),
}


def extract_template_placeholders(text: str) -> List[str]:
    """Return normalized placeholder names found in a template text preview."""
    found: List[str] = []
    for match in PLACEHOLDER_PATTERN.finditer(text or ""):
        raw = match.group(1) or match.group(2) or ""
        key = normalize_placeholder_key(raw)
        if key and key not in found:
            found.append(key)
    return found


def normalize_placeholder_key(value: str) -> str:
    raw = re.sub(r"\s+", "", str(value or "").strip().lower())
    raw = raw.replace("-", "_")
    raw = raw.replace("：", "").replace(":", "")
    return KEY_ALIASES.get(raw, raw)


def make_template_preview(text: str, *, limit: int = 900) -> str:
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    return cleaned[:limit]


def build_lesson_plan_docx(
    plan_data: Dict[str, Any],
    meta: Dict[str, Any],
    out_path: Path,
    *,
    template_path: Optional[Path] = None,
    template_text: str = "",
) -> None:
    """Build a DOCX lesson plan, filling placeholders or applying template styles."""
    from docx import Document

    doc = None
    if template_path and template_path.exists() and template_path.suffix.lower() == ".docx":
        try:
            doc = Document(str(template_path))
        except Exception:
            doc = None
    if doc is None:
        doc = Document()
        _apply_default_document_style(doc)

    values = _template_values(plan_data, meta)
    replaced_count = 0
    replaced_keys: Set[str] = set()

    if template_path and template_path.exists() and template_path.suffix.lower() == ".docx":
        replaced_count, replaced_keys = _replace_document_placeholders(doc, values)
        auto_keys = _auto_fill_document_template(doc, values)
        mapped_keys = replaced_keys | auto_keys
        if not (mapped_keys & CORE_CONTENT_KEYS):
            _append_standard_lesson_plan(doc, plan_data, include_title=False)
    elif template_text:
        mapped_keys = _build_from_text_template(doc, template_text, values)
        if not (mapped_keys & CORE_CONTENT_KEYS):
            _append_standard_lesson_plan(doc, plan_data, include_title=False)
    else:
        _clear_document_body(doc)
        _append_standard_lesson_plan(doc, plan_data, include_title=True)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _apply_default_document_style(doc: Any) -> None:
    try:
        from docx.shared import Pt

        normal = doc.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(11)
    except Exception:
        pass


def _template_values(plan_data: Dict[str, Any], meta: Dict[str, Any]) -> Dict[str, Any]:
    goals = plan_data.get("teaching_goals") if isinstance(plan_data.get("teaching_goals"), dict) else {}
    key_points = _string_list(plan_data.get("key_points"))
    difficult_points = _string_list(plan_data.get("difficult_points"))
    preparation = _string_list(plan_data.get("preparation"))
    process = _process_items(plan_data.get("process"))

    knowledge = str(goals.get("knowledge") or "").strip()
    process_goal = str(goals.get("process") or "").strip()
    emotion = str(goals.get("emotion") or "").strip()

    return {
        "title": str(plan_data.get("title") or meta.get("title") or "教案").strip(),
        "generated_at": str(meta.get("created_at") or meta.get("updated_at") or "").strip(),
        "teaching_goals": "\n".join(
            item
            for item in [
                f"知识与技能：{knowledge}" if knowledge else "",
                f"过程与方法：{process_goal}" if process_goal else "",
                f"情感态度与价值观：{emotion}" if emotion else "",
            ]
            if item
        ),
        "knowledge": knowledge,
        "process_goal": process_goal,
        "emotion": emotion,
        "key_points": _numbered_lines(key_points),
        "difficult_points": _numbered_lines(difficult_points),
        "key_and_difficult_points": "\n".join(
            item
            for item in [
                f"教学重点：{'；'.join(key_points)}" if key_points else "",
                f"教学难点：{'；'.join(difficult_points)}" if difficult_points else "",
            ]
            if item
        ),
        "preparation": _numbered_lines(preparation),
        "process": "\n\n".join(
            f"{item['title']}\n{item['content']}".strip()
            for item in process
            if item.get("title") or item.get("content")
        ),
        "_process_items": process,
        "homework": str(plan_data.get("homework") or "").strip(),
        "blackboard_design": _blackboard_design(plan_data, key_points, difficult_points),
    }


def _replace_document_placeholders(doc: Any, values: Dict[str, Any]) -> Tuple[int, Set[str]]:
    count = 0
    keys: Set[str] = set()
    for paragraph in _iter_all_paragraphs(doc):
        changed, changed_keys = _replace_paragraph_placeholders(paragraph, values)
        if changed:
            count += 1
            keys.update(changed_keys)
    return count, keys


def _iter_all_paragraphs(doc: Any) -> Iterable[Any]:
    yield from _iter_container_paragraphs(doc)
    for section in doc.sections:
        yield from _iter_container_paragraphs(section.header)
        yield from _iter_container_paragraphs(section.footer)


def _iter_container_paragraphs(container: Any) -> Iterable[Any]:
    for paragraph in getattr(container, "paragraphs", []):
        yield paragraph
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_container_paragraphs(cell)


def _replace_paragraph_placeholders(paragraph: Any, values: Dict[str, Any]) -> Tuple[bool, Set[str]]:
    original = paragraph.text
    changed_keys: Set[str] = set()

    def replace(match: re.Match[str]) -> str:
        raw = match.group(1) or match.group(2) or ""
        key = normalize_placeholder_key(raw)
        if key in values:
            changed_keys.add(key)
            return str(values[key])
        return match.group(0)

    updated = PLACEHOLDER_PATTERN.sub(replace, original)
    if updated == original:
        return False, changed_keys
    _set_paragraph_text(paragraph, updated)
    return True, changed_keys


def _replace_text_placeholders(text: str, values: Dict[str, Any]) -> str:
    def replace(match: re.Match[str]) -> str:
        key = normalize_placeholder_key(match.group(1) or match.group(2) or "")
        return str(values.get(key, match.group(0)))

    return PLACEHOLDER_PATTERN.sub(replace, text or "")


def _auto_fill_document_template(doc: Any, values: Dict[str, Any]) -> Set[str]:
    """Fill common no-placeholder Word templates in place."""
    filled: Set[str] = set()
    filled.update(_fill_table_templates(doc, values))
    filled.update(_fill_paragraph_templates(doc, values))
    return filled


def _fill_table_templates(doc: Any, values: Dict[str, Any]) -> Set[str]:
    filled: Set[str] = set()
    for table in _iter_all_tables(doc):
        filled.update(_fill_single_table(table, values))
    return filled


def _iter_all_tables(doc: Any) -> Iterable[Any]:
    yield from _iter_container_tables(doc)
    for section in doc.sections:
        yield from _iter_container_tables(section.header)
        yield from _iter_container_tables(section.footer)


def _iter_container_tables(container: Any) -> Iterable[Any]:
    for table in getattr(container, "tables", []):
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_container_tables(cell)


def _fill_single_table(table: Any, values: Dict[str, Any]) -> Set[str]:
    filled: Set[str] = set()
    if _fill_process_grid_table(table, values):
        return {"process"}
    rows = list(table.rows)
    for row_index, row in enumerate(rows):
        cells = list(row.cells)
        seen_cells: Set[int] = set()
        for cell_index, cell in enumerate(cells):
            tc_id = id(cell._tc)
            if tc_id in seen_cells:
                continue
            seen_cells.add(tc_id)
            text = _cell_text(cell)
            key = _detect_field_key(text)
            value = str(values.get(key or "") or "").strip()
            if not key or not value:
                continue

            target = _find_row_target_cell(cells, cell_index + 1)
            if target is None:
                target = _find_below_target_cell(rows, row_index, cell_index)
            if target is not None:
                _set_cell_text(target, value)
                filled.add(key)
                continue

            if _can_fill_inline_label(text, key):
                _set_cell_text(cell, _inline_filled_text(text, key, value))
                filled.add(key)
            elif _is_label_only_text(text, key):
                _set_cell_text(cell, f"{text.strip()}\n{value}")
                filled.add(key)
    return filled


def _fill_process_grid_table(table: Any, values: Dict[str, Any]) -> bool:
    """Fill teaching-process grid tables such as 环节/教师活动/学生活动/设计意图."""
    rows = list(table.rows)
    header_index, column_roles = _find_process_grid_header(rows)
    if header_index < 0:
        return False

    process_items = values.get("_process_items")
    if not isinstance(process_items, list):
        process_items = []
    homework = str(values.get("homework") or "").strip()
    if not process_items and not homework:
        return False

    phase_col = _first_column_for_role(column_roles, "phase")
    if phase_col is None:
        phase_col = 0

    wrote = False
    for row_offset, row in enumerate(rows[header_index + 1 :], start=0):
        cells = list(row.cells)
        if not cells:
            continue

        phase_cell = cells[phase_col] if phase_col < len(cells) else cells[0]
        phase_text = _cell_text(phase_cell)
        if _row_looks_like_process_header(cells):
            continue

        phase_kind = _process_phase_kind(phase_text)
        item = _select_process_item_for_phase(phase_text, process_items, row_offset)
        if not phase_kind:
            phase_kind = _process_item_kind(item) if item else ""
        if not phase_kind and not item:
            continue

        written_cell_ids: Set[int] = set()
        for col_index, role in column_roles.items():
            if role == "phase" or col_index >= len(cells):
                continue
            cell = cells[col_index]
            tc_id = id(cell._tc)
            if tc_id in written_cell_ids or tc_id == id(phase_cell._tc):
                continue
            written_cell_ids.add(tc_id)

            text = _process_grid_value_for_role(role, phase_kind, item, homework)
            if not text:
                continue

            current = _cell_text(cell)
            if role == "time":
                if _is_fillable_text(current):
                    _set_cell_text(cell, text)
                    wrote = True
                continue

            if _should_fill_process_grid_cell(current):
                _set_cell_text(cell, text)
                wrote = True

    return wrote


def _find_process_grid_header(rows: Sequence[Any]) -> Tuple[int, Dict[int, str]]:
    for row_index, row in enumerate(rows[:6]):
        roles: Dict[int, str] = {}
        seen_cells: Set[int] = set()
        for col_index, cell in enumerate(row.cells):
            tc_id = id(cell._tc)
            if tc_id in seen_cells:
                continue
            seen_cells.add(tc_id)
            role = _detect_process_column_role(_cell_text(cell))
            if role:
                roles[col_index] = role

        role_values = set(roles.values())
        detail_count = len(role_values & PROCESS_GRID_DETAIL_ROLES)
        if ("phase" in role_values and detail_count >= 1) or detail_count >= 3:
            if "phase" not in role_values:
                roles.setdefault(0, "phase")
            return row_index, roles
    return -1, {}


def _detect_process_column_role(text: str) -> str:
    normalized = _normalize_label_text(text)
    if not normalized:
        return ""
    if any(word in normalized for word in ("教学环节", "环节", "步骤", "流程")):
        return "phase"
    if any(word in normalized for word in ("时间分配", "时间安排", "用时", "时间")):
        return "time"
    if any(word in normalized for word in ("教师活动", "教师行为", "师生活动", "教师")):
        return "teacher"
    if any(word in normalized for word in ("学生活动", "学生行为", "学习活动", "学生")):
        return "student"
    if any(word in normalized for word in ("设计意图", "设计目的", "设计说明", "意图")):
        return "intent"
    if any(word in normalized for word in ("教学内容", "活动内容", "主要内容", "教学任务")):
        return "content"
    return ""


def _first_column_for_role(column_roles: Dict[int, str], role: str) -> Optional[int]:
    for col_index, value in column_roles.items():
        if value == role:
            return col_index
    return None


def _row_looks_like_process_header(cells: Sequence[Any]) -> bool:
    roles = {_detect_process_column_role(_cell_text(cell)) for cell in cells}
    roles.discard("")
    return "phase" in roles and bool(roles & PROCESS_GRID_DETAIL_ROLES)


def _select_process_item_for_phase(
    phase_text: str,
    process_items: Sequence[Dict[str, str]],
    row_offset: int,
) -> Optional[Dict[str, str]]:
    if not process_items:
        return None

    phase_kind = _process_phase_kind(phase_text)
    if phase_kind == "homework":
        return None

    if phase_kind:
        for item in process_items:
            if _process_item_kind(item) == phase_kind:
                return item
        if phase_kind == "explore":
            for item in process_items:
                if _process_item_kind(item) == "new":
                    return item
        fallback_index = _fallback_process_index(phase_kind, len(process_items))
        if fallback_index is not None:
            return process_items[fallback_index]

    if row_offset < len(process_items):
        return process_items[row_offset]
    return process_items[-1]


def _fallback_process_index(phase_kind: str, count: int) -> Optional[int]:
    if count <= 0:
        return None
    if phase_kind == "lead_in":
        return 0
    if phase_kind == "new":
        return 1 if count > 1 else 0
    if phase_kind == "explore":
        return 2 if count >= 5 else (1 if count > 1 else 0)
    if phase_kind == "practice":
        return 3 if count >= 5 else (2 if count > 2 else count - 1)
    if phase_kind == "summary":
        return count - 1
    return None


def _process_item_kind(item: Optional[Dict[str, str]]) -> str:
    if not item:
        return ""
    title_kind = _process_phase_kind(str(item.get("title") or ""))
    if title_kind:
        return title_kind
    return _process_phase_kind(str(item.get("content") or "")[:100])


def _process_phase_kind(text: str) -> str:
    normalized = _normalize_label_text(text)
    if not normalized:
        return ""
    for phase_kind, keywords in PROCESS_PHASE_KEYWORDS.items():
        if any(keyword in normalized for keyword in keywords):
            return phase_kind
    return ""


def _process_grid_value_for_role(
    role: str,
    phase_kind: str,
    item: Optional[Dict[str, str]],
    homework: str,
) -> str:
    if role == "time":
        return _process_time_text(phase_kind)

    if phase_kind == "homework":
        if role in {"teacher", "content"}:
            return homework or "布置分层作业，明确完成要求与提交方式。"
        if role == "student":
            return "记录作业要求，明确完成标准，课后独立完成。"
        if role == "intent":
            return "延伸课堂学习，巩固重点内容，促进迁移应用。"

    title = _strip_process_title(str((item or {}).get("title") or ""))
    content = str((item or {}).get("content") or "").strip()
    if role == "teacher":
        return content or _generic_teacher_activity(phase_kind)
    if role == "student":
        return _generic_student_activity(phase_kind)
    if role == "intent":
        return _generic_design_intent(phase_kind)
    if role == "content":
        return f"{title}\n{content}".strip() if title else content
    return ""


def _process_time_text(phase_kind: str) -> str:
    return {
        "lead_in": "5分钟",
        "new": "15分钟",
        "explore": "10分钟",
        "practice": "8分钟",
        "summary": "5分钟",
        "homework": "2分钟",
    }.get(phase_kind, "")


def _generic_teacher_activity(phase_kind: str) -> str:
    return {
        "lead_in": "创设情境，提出问题，引导学生进入本课学习任务。",
        "new": "围绕核心知识分步讲解，组织提问与示范，突破重点难点。",
        "explore": "设计探究任务，巡视指导小组合作，组织交流展示。",
        "practice": "布置分层练习，及时反馈共性问题，指导学生订正。",
        "summary": "引导学生梳理知识结构，总结方法，回应课堂疑问。",
    }.get(phase_kind, "组织学习活动，引导学生围绕本课目标完成任务。")


def _generic_student_activity(phase_kind: str) -> str:
    return {
        "lead_in": "观察情境，回答问题，明确本课学习任务。",
        "new": "跟随示范思考，记录关键概念和方法，参与课堂问答。",
        "explore": "小组合作交流，提出猜想，展示探究过程和结论。",
        "practice": "独立完成练习，交流解题思路，订正错误。",
        "summary": "归纳本课要点，交流收获与疑问。",
    }.get(phase_kind, "参与讨论，记录要点，完成对应学习任务。")


def _generic_design_intent(phase_kind: str) -> str:
    return {
        "lead_in": "激活已有经验，形成学习期待，建立新旧知识联系。",
        "new": "突出教学重点，分步突破难点，帮助学生建构知识。",
        "explore": "通过合作交流促进深度理解，提升探究与表达能力。",
        "practice": "检测学习效果，巩固关键知识与技能。",
        "summary": "完善知识结构，促进反思内化。",
    }.get(phase_kind, "服务本课教学目标，促进学生理解与迁移。")


def _should_fill_process_grid_cell(text: str) -> bool:
    if _is_fillable_text(text):
        return True
    if _detect_process_column_role(text) or _detect_field_key(text):
        return False
    compact = _normalize_label_text(text)
    return bool(compact) and len(compact) <= 120


def _strip_process_title(text: str) -> str:
    return re.sub(r"^\s*[\d一二三四五六七八九十]+[、.．)\）-]*\s*", "", str(text or "")).strip()


def _find_row_target_cell(cells: Sequence[Any], start_index: int) -> Optional[Any]:
    seen_cells: Set[int] = set()
    for cell in cells[start_index:]:
        tc_id = id(cell._tc)
        if tc_id in seen_cells:
            continue
        seen_cells.add(tc_id)
        text = _cell_text(cell)
        if _detect_field_key(text):
            break
        if _is_fillable_text(text):
            return cell
    return None


def _find_below_target_cell(rows: Sequence[Any], row_index: int, cell_index: int) -> Optional[Any]:
    for next_row in rows[row_index + 1 : row_index + 4]:
        cells = list(next_row.cells)
        unique_cells = []
        unique_ids: Set[int] = set()
        for cell in cells:
            tc_id = id(cell._tc)
            if tc_id not in unique_ids:
                unique_ids.add(tc_id)
                unique_cells.append(cell)
        if any(_detect_field_key(_cell_text(cell)) for cell in unique_cells):
            continue
        ordered = cells[cell_index:] + cells[:cell_index]
        seen_cells: Set[int] = set()
        for cell in ordered:
            tc_id = id(cell._tc)
            if tc_id in seen_cells:
                continue
            seen_cells.add(tc_id)
            text = _cell_text(cell)
            if _detect_field_key(text):
                continue
            if _is_fillable_text(text):
                return cell
    return None


def _fill_paragraph_templates(doc: Any, values: Dict[str, Any]) -> Set[str]:
    filled: Set[str] = set()
    paragraphs = list(getattr(doc, "paragraphs", []))
    original_texts = [paragraph.text for paragraph in paragraphs]
    detail_goal_indices = {
        index
        for index, text in enumerate(original_texts)
        if _detect_field_key(text) in {"knowledge", "process_goal", "emotion"}
    }
    for index, paragraph in enumerate(paragraphs):
        text = original_texts[index]
        key = _detect_field_key(text)
        value = str(values.get(key or "") or "").strip()
        if not key or not value:
            continue
        if key == "teaching_goals" and any(index < detail_index <= index + 5 for detail_index in detail_goal_indices):
            continue

        if _can_fill_inline_label(text, key):
            _set_paragraph_text(paragraph, _inline_filled_text(text, key, value))
            filled.add(key)
            continue

        next_paragraph = paragraphs[index + 1] if index + 1 < len(paragraphs) else None
        next_text = original_texts[index + 1] if index + 1 < len(paragraphs) else ""
        if next_paragraph is not None and not _detect_field_key(next_text) and _is_fillable_text(next_text):
            _set_paragraph_text(next_paragraph, value)
            filled.add(key)
        elif _is_label_only_text(text, key):
            _insert_paragraph_after(paragraph, value)
            filled.add(key)
    return filled


def _build_from_text_template(doc: Any, template_text: str, values: Dict[str, Any]) -> Set[str]:
    _clear_document_body(doc)
    filled: Set[str] = set()
    for raw_line in str(template_text or "").splitlines():
        line = raw_line.strip()
        placeholder_keys = set(extract_template_placeholders(line))
        if placeholder_keys:
            doc.add_paragraph(_replace_text_placeholders(line, values))
            filled.update(key for key in placeholder_keys if key in values)
            continue
        key = _detect_field_key(line)
        value = str(values.get(key or "") or "").strip()
        if key and value:
            if _can_fill_inline_label(line, key):
                doc.add_paragraph(_inline_filled_text(line, key, value))
            else:
                if line:
                    doc.add_paragraph(line)
                doc.add_paragraph(value)
            filled.add(key)
        else:
            doc.add_paragraph(line)
    return filled


def _detect_field_key(text: str) -> str:
    normalized = _normalize_label_text(text)
    if not normalized:
        return ""
    direct = KEY_ALIASES.get(normalized)
    if direct:
        return direct
    for key, labels in LABEL_PRIORITY:
        for label in labels:
            label_norm = _normalize_label_text(label)
            if label_norm in normalized and _looks_like_template_label(text, normalized, label_norm):
                return key
    return ""


def _normalize_label_text(text: str) -> str:
    value = re.sub(r"\s+", "", str(text or "").strip().lower())
    value = re.sub(r"^[一二三四五六七八九十\d]+[、.．)）-]*", "", value)
    value = re.sub(r"[：:;；,，。.!！?？（）()\[\]【】《》<>_＿\-—–\s]+", "", value)
    return value


def _looks_like_template_label(original: str, normalized: str, label_norm: str) -> bool:
    if not normalized or not label_norm:
        return False
    if len(normalized) <= len(label_norm) + 18:
        return True
    if _has_fill_marker(original):
        return True
    first = normalized.find(label_norm)
    return first <= 2 and len(normalized) <= len(label_norm) + 28


def _cell_text(cell: Any) -> str:
    return "\n".join(paragraph.text for paragraph in getattr(cell, "paragraphs", []))


def _set_cell_text(cell: Any, text: str) -> None:
    paragraphs = list(getattr(cell, "paragraphs", []))
    if not paragraphs:
        cell.add_paragraph()
        paragraphs = list(cell.paragraphs)
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    _set_paragraph_text(paragraphs[0], text)


def _is_fillable_text(text: str) -> bool:
    compact = re.sub(r"\s+", "", str(text or ""))
    if not compact:
        return True
    if PLACEHOLDER_PATTERN.search(text):
        return True
    if re.fullmatch(r"[_＿\-—–·.。…]+", compact):
        return True
    if len(compact) <= 50 and len(re.findall(r"[\u4e00-\u9fff]\s+[\u4e00-\u9fff]", str(text or ""))) >= 2:
        return True
    fill_words = ("请填写", "待填写", "填写", "空白", "无", "暂无", "待补充")
    if any(word in compact for word in fill_words):
        return True
    return False


def _can_fill_inline_label(text: str, key: str) -> bool:
    if not text or not _detect_field_key(text) == key:
        return False
    if PLACEHOLDER_PATTERN.search(text):
        return True
    if _has_fill_marker(text):
        return True
    if re.search(r"[：:]\s*(请填写|待填写|填写|空白|无|暂无|待补充)?\s*$", text):
        return True
    return False


def _has_fill_marker(text: str) -> bool:
    return bool(re.search(r"[_＿]{2,}|[—-]{2,}|…{1,}", str(text or "")))


def _is_label_only_text(text: str, key: str) -> bool:
    if not text or _detect_field_key(text) != key:
        return False
    normalized = _normalize_label_text(text)
    label_norms = [_normalize_label_text(label) for item in LABEL_PRIORITY if item[0] == key for label in item[1]]
    if normalized in label_norms:
        return True
    label_suffixes = ("目标", "内容", "设计", "安排", "要求")
    return any(normalized == f"{label}{suffix}" for label in label_norms for suffix in label_suffixes)


def _inline_filled_text(text: str, key: str, value: str) -> str:
    original = str(text or "").strip()
    if PLACEHOLDER_PATTERN.search(original):
        return _replace_text_placeholders(original, {key: value})
    colon_match = re.search(r"^(.*?[：:])", original)
    if colon_match:
        return f"{colon_match.group(1)}{value}"
    blank_match = re.search(r"^(.*?)(?:[_＿]{2,}|[—-]{2,}|…{1,})", original)
    if blank_match:
        return f"{blank_match.group(1)}{value}"
    return f"{original}\n{value}"


def _insert_paragraph_after(paragraph: Any, text: str) -> Any:
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    try:
        new_paragraph.style = paragraph._parent.styles["Normal"]
    except Exception:
        pass
    _set_paragraph_text(new_paragraph, text)
    return new_paragraph


def _set_paragraph_text(paragraph: Any, text: str) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run()
    parts = str(text or "").split("\n")
    for index, part in enumerate(parts):
        if index:
            run.add_break()
        run.add_text(part)


def _clear_document_body(doc: Any) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _append_text_template(doc: Any, text: str) -> None:
    for line in str(text or "").splitlines():
        doc.add_paragraph(line.strip() if line.strip() else "")


def _append_standard_lesson_plan(doc: Any, plan_data: Dict[str, Any], *, include_title: bool) -> None:
    title = str(plan_data.get("title") or "教案").strip()
    if include_title:
        _add_title(doc, title)

    goals = plan_data.get("teaching_goals") if isinstance(plan_data.get("teaching_goals"), dict) else {}
    _add_heading(doc, "一、教学目标（三维目标）")
    if goals.get("knowledge"):
        _add_label_paragraph(doc, "知识与技能：", str(goals.get("knowledge") or ""))
    if goals.get("process"):
        _add_label_paragraph(doc, "过程与方法：", str(goals.get("process") or ""))
    if goals.get("emotion"):
        _add_label_paragraph(doc, "情感态度与价值观：", str(goals.get("emotion") or ""))

    _add_heading(doc, "二、教学重难点")
    _add_items_or_empty(doc, "教学重点", _string_list(plan_data.get("key_points")))
    _add_items_or_empty(doc, "教学难点", _string_list(plan_data.get("difficult_points")))

    _add_heading(doc, "三、教学准备")
    _add_bullets(doc, _string_list(plan_data.get("preparation")) or ["—"])

    _add_heading(doc, "四、教学过程")
    for item in _process_items(plan_data.get("process")):
        _add_subheading(doc, item.get("title") or "环节")
        doc.add_paragraph(item.get("content") or "")

    _add_heading(doc, "五、作业设计")
    doc.add_paragraph(str(plan_data.get("homework") or "—"))


def _add_title(doc: Any, text: str) -> None:
    if _style_exists(doc, "Title"):
        paragraph = doc.add_paragraph(text, style="Title")
    else:
        paragraph = doc.add_paragraph(text)
    try:
        paragraph.alignment = 1
        if paragraph.runs:
            paragraph.runs[0].bold = True
    except Exception:
        pass


def _add_heading(doc: Any, text: str) -> None:
    if _style_exists(doc, "Heading 1"):
        doc.add_heading(text, level=1)
    else:
        paragraph = doc.add_paragraph(text)
        if paragraph.runs:
            paragraph.runs[0].bold = True


def _add_subheading(doc: Any, text: str) -> None:
    if _style_exists(doc, "Heading 2"):
        doc.add_heading(text, level=2)
    else:
        paragraph = doc.add_paragraph(text)
        if paragraph.runs:
            paragraph.runs[0].bold = True


def _add_label_paragraph(doc: Any, label: str, content: str) -> None:
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(label)
    label_run.bold = True
    paragraph.add_run(content)


def _add_items_or_empty(doc: Any, label: str, items: Sequence[str]) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{label}：")
    run.bold = True
    paragraph.add_run("；".join(items) if items else "—")


def _add_bullets(doc: Any, items: Sequence[str]) -> None:
    for item in items:
        if _style_exists(doc, "List Bullet"):
            doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph(str(item))


def _style_exists(doc: Any, style_name: str) -> bool:
    try:
        doc.styles[style_name]
        return True
    except Exception:
        return False


def _string_list(value: Any) -> List[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item or "").strip()]


def _process_items(value: Any) -> List[Dict[str, str]]:
    result: List[Dict[str, str]] = []
    if isinstance(value, list):
        for item in value:
            if isinstance(item, dict):
                result.append(
                    {
                        "title": str(item.get("title") or "").strip(),
                        "content": str(item.get("content") or "").strip(),
                    }
                )
            elif str(item or "").strip():
                result.append({"title": "环节", "content": str(item).strip()})
    return result


def _numbered_lines(items: Sequence[str]) -> str:
    return "\n".join(f"{index}. {item}" for index, item in enumerate(items, start=1))


def _blackboard_design(plan_data: Dict[str, Any], key_points: Sequence[str], difficult_points: Sequence[str]) -> str:
    title = str(plan_data.get("title") or "本课").strip()
    lines = [title]
    if key_points:
        lines.append("一、教学重点")
        lines.extend(f"{index}. {item}" for index, item in enumerate(key_points, start=1))
    if difficult_points:
        lines.append("二、难点突破")
        lines.extend(f"{index}. {item}" for index, item in enumerate(difficult_points, start=1))
    if len(lines) == 1:
        lines.append("围绕核心概念、关键问题和课堂小结形成板书。")
    return "\n".join(lines)
